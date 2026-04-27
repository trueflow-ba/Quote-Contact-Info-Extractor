from dotenv import load_dotenv
from pathlib import Path

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

import os
import uuid
import io
import csv
import json
import zipfile
import asyncio
import logging
import re
from datetime import datetime, timezone, timedelta
from typing import List, Optional

from fastapi import FastAPI, APIRouter, Request, Response, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
from pydantic import BaseModel
from bson import ObjectId
import bcrypt
import jwt as pyjwt
import requests as http_requests
import pdfplumber

from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContent, ImageContent

# =============================================================================
# CONFIG
# =============================================================================
JWT_ALGORITHM = "HS256"
STORAGE_URL = "https://integrations.emergentagent.com/objstore/api/v1/storage"
APP_NAME = "trueflow"
BATCH_SIZE = 10

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI()
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# =============================================================================
# STORAGE HELPERS
# =============================================================================
storage_key = None

def init_storage():
    global storage_key
    if storage_key:
        return storage_key
    emergent_key = os.environ.get("EMERGENT_LLM_KEY")
    if not emergent_key:
        logger.warning("EMERGENT_LLM_KEY not set, storage disabled")
        return None
    resp = http_requests.post(f"{STORAGE_URL}/init", json={"emergent_key": emergent_key}, timeout=30)
    resp.raise_for_status()
    storage_key = resp.json()["storage_key"]
    return storage_key

def put_object(path: str, data: bytes, content_type: str) -> dict:
    key = init_storage()
    if not key:
        raise HTTPException(status_code=500, detail="Storage not initialized")
    last_err = None
    for attempt in range(3):
        try:
            resp = http_requests.put(
                f"{STORAGE_URL}/objects/{path}",
                headers={"X-Storage-Key": key, "Content-Type": content_type},
                data=data, timeout=120
            )
            resp.raise_for_status()
            return resp.json()
        except Exception as e:
            last_err = e
            if attempt < 2:
                import time
                time.sleep(2 ** attempt)
                # Re-init storage key in case it expired
                global storage_key
                storage_key = None
                try:
                    init_storage()
                except:
                    pass
    raise last_err

def get_object(path: str) -> bytes:
    key = init_storage()
    if not key:
        raise HTTPException(status_code=500, detail="Storage not initialized")
    resp = http_requests.get(
        f"{STORAGE_URL}/objects/{path}",
        headers={"X-Storage-Key": key}, timeout=60
    )
    resp.raise_for_status()
    return resp.content

def delete_object(path: str):
    key = init_storage()
    if not key:
        return
    try:
        http_requests.delete(
            f"{STORAGE_URL}/objects/{path}",
            headers={"X-Storage-Key": key}, timeout=30
        )
    except Exception as e:
        logger.warning(f"Failed to delete storage object {path}: {e}")

# =============================================================================
# AUTH HELPERS
# =============================================================================
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")

def verify_password(plain: str, hashed: str) -> bool:
    return bcrypt.checkpw(plain.encode("utf-8"), hashed.encode("utf-8"))

def get_jwt_secret():
    return os.environ["JWT_SECRET"]

def create_access_token(user_id: str, email: str) -> str:
    payload = {"sub": user_id, "email": email, "exp": datetime.now(timezone.utc) + timedelta(minutes=15), "type": "access"}
    return pyjwt.encode(payload, get_jwt_secret(), algorithm=JWT_ALGORITHM)

def create_refresh_token(user_id: str) -> str:
    payload = {"sub": user_id, "exp": datetime.now(timezone.utc) + timedelta(days=7), "type": "refresh"}
    return pyjwt.encode(payload, get_jwt_secret(), algorithm=JWT_ALGORITHM)

async def get_current_user(request: Request) -> dict:
    token = request.cookies.get("access_token")
    if not token:
        auth_header = request.headers.get("Authorization", "")
        if auth_header.startswith("Bearer "):
            token = auth_header[7:]
    if not token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = pyjwt.decode(token, get_jwt_secret(), algorithms=[JWT_ALGORITHM])
        if payload.get("type") != "access":
            raise HTTPException(status_code=401, detail="Invalid token type")
        user = await db.users.find_one({"_id": ObjectId(payload["sub"])})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        user["_id"] = str(user["_id"])
        user.pop("password_hash", None)
        return user
    except pyjwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except pyjwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")

# =============================================================================
# PYDANTIC MODELS
# =============================================================================
class RegisterInput(BaseModel):
    email: str
    password: str
    name: str = ""

class LoginInput(BaseModel):
    email: str
    password: str

class SettingsInput(BaseModel):
    exclusion_domain: Optional[str] = None

class AdminSettingsInput(BaseModel):
    ai_model: Optional[str] = None
    claude_api_key: Optional[str] = None
    openai_api_key: Optional[str] = None
    max_pdfs_per_upload: Optional[int] = None
    storage_max_mb: Optional[int] = None
    storage_target_mb: Optional[int] = None
    # P0 prework: safety controls for large runs
    budget_ceiling_usd: Optional[float] = None           # Auto-pause when approx_cost exceeds this
    consecutive_failure_threshold: Optional[int] = None  # Auto-pause after N consecutive file failures
    retry_max_attempts: Optional[int] = None             # Per-file LLM call retry attempts

class ChangePasswordInput(BaseModel):
    current_password: str
    new_password: str

# =============================================================================
# AUTH ROUTES
# =============================================================================
@api_router.post("/auth/register")
async def register(input: RegisterInput, response: Response):
    email = input.email.lower().strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Invalid email address")
    if len(input.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")
    existing = await db.users.find_one({"email": email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    user_doc = {
        "email": email,
        "password_hash": hash_password(input.password),
        "name": input.name or email.split("@")[0],
        "role": "user",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    result = await db.users.insert_one(user_doc)
    user_id = str(result.inserted_id)
    await db.settings.insert_one({
        "user_id": user_id,
        "ai_model": "claude-sonnet",
        "claude_api_key": "",
        "openai_api_key": "",
        "exclusion_domain": "horizonc.com",
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    access_token = create_access_token(user_id, email)
    refresh_token = create_refresh_token(user_id)
    response.set_cookie(key="access_token", value=access_token, httponly=True, secure=False, samesite="lax", max_age=900, path="/")
    response.set_cookie(key="refresh_token", value=refresh_token, httponly=True, secure=False, samesite="lax", max_age=604800, path="/")
    return {"_id": user_id, "email": email, "name": user_doc["name"], "role": "user"}

@api_router.post("/auth/login")
async def login(input: LoginInput, request: Request, response: Response):
    email = input.email.lower().strip()
    ip = request.client.host if request.client else "unknown"
    identifier = f"{ip}:{email}"
    attempt = await db.login_attempts.find_one({"identifier": identifier})
    if attempt and attempt.get("attempts", 0) >= 5:
        locked_until = attempt.get("locked_until")
        if locked_until and datetime.now(timezone.utc).isoformat() < locked_until:
            raise HTTPException(status_code=429, detail="Too many failed attempts. Try again in 15 minutes.")
        else:
            await db.login_attempts.delete_one({"identifier": identifier})
    user = await db.users.find_one({"email": email})
    if not user or not verify_password(input.password, user["password_hash"]):
        await db.login_attempts.update_one(
            {"identifier": identifier},
            {"$inc": {"attempts": 1}, "$set": {"locked_until": (datetime.now(timezone.utc) + timedelta(minutes=15)).isoformat()}},
            upsert=True
        )
        raise HTTPException(status_code=401, detail="Invalid email or password")
    await db.login_attempts.delete_one({"identifier": identifier})
    user_id = str(user["_id"])
    access_token = create_access_token(user_id, email)
    refresh_token = create_refresh_token(user_id)
    response.set_cookie(key="access_token", value=access_token, httponly=True, secure=False, samesite="lax", max_age=900, path="/")
    response.set_cookie(key="refresh_token", value=refresh_token, httponly=True, secure=False, samesite="lax", max_age=604800, path="/")
    return {"_id": user_id, "email": email, "name": user.get("name", ""), "role": user.get("role", "user"), "must_change_password": user.get("must_change_password", False)}

@api_router.post("/auth/logout")
async def logout(response: Response):
    response.delete_cookie("access_token", path="/")
    response.delete_cookie("refresh_token", path="/")
    return {"message": "Logged out"}

@api_router.get("/auth/me")
async def get_me(request: Request):
    return await get_current_user(request)

@api_router.post("/auth/refresh")
async def refresh(request: Request, response: Response):
    token = request.cookies.get("refresh_token")
    if not token:
        raise HTTPException(status_code=401, detail="No refresh token")
    try:
        payload = pyjwt.decode(token, get_jwt_secret(), algorithms=[JWT_ALGORITHM])
        if payload.get("type") != "refresh":
            raise HTTPException(status_code=401, detail="Invalid token type")
        user = await db.users.find_one({"_id": ObjectId(payload["sub"])})
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        user_id = str(user["_id"])
        access_token = create_access_token(user_id, user["email"])
        response.set_cookie(key="access_token", value=access_token, httponly=True, secure=False, samesite="lax", max_age=900, path="/")
        return {"message": "Token refreshed"}
    except pyjwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Refresh token expired")
    except pyjwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid refresh token")

# =============================================================================
# AUTH - CHANGE PASSWORD
# =============================================================================
@api_router.post("/auth/change-password")
async def change_password(input: ChangePasswordInput, request: Request):
    user_full = await get_current_user(request)
    user_doc = await db.users.find_one({"_id": ObjectId(user_full["_id"])})
    if not user_doc:
        raise HTTPException(status_code=404, detail="User not found")
    if not verify_password(input.current_password, user_doc["password_hash"]):
        raise HTTPException(status_code=400, detail="Current password is incorrect")
    if len(input.new_password) < 6:
        raise HTTPException(status_code=400, detail="New password must be at least 6 characters")
    await db.users.update_one(
        {"_id": ObjectId(user_full["_id"])},
        {"$set": {"password_hash": hash_password(input.new_password), "must_change_password": False}}
    )
    return {"message": "Password changed successfully"}

# =============================================================================
# USER SETTINGS (exclusion domain only - available to all users)
# =============================================================================
@api_router.get("/settings")
async def get_settings(request: Request):
    user = await get_current_user(request)
    settings = await db.settings.find_one({"user_id": user["_id"]}, {"_id": 0})
    if not settings:
        settings = {"user_id": user["_id"], "exclusion_domain": "horizonc.com"}
    # Also fetch admin config so user knows the model and limits
    admin_config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
    settings["ai_model"] = admin_config.get("ai_model", "claude-sonnet") if admin_config else "claude-sonnet"
    settings["max_pdfs_per_upload"] = admin_config.get("max_pdfs_per_upload", 50) if admin_config else 50
    return settings

@api_router.put("/settings")
async def update_settings(input: SettingsInput, request: Request):
    user = await get_current_user(request)
    update = {"updated_at": datetime.now(timezone.utc).isoformat()}
    if input.exclusion_domain is not None:
        update["exclusion_domain"] = input.exclusion_domain
    await db.settings.update_one({"user_id": user["_id"]}, {"$set": update}, upsert=True)
    return {"message": "Settings updated"}

# =============================================================================
# ADMIN SETTINGS (AI model, API keys, PDF limits - admin only)
# =============================================================================
async def require_admin(request: Request) -> dict:
    user = await get_current_user(request)
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return user

@api_router.get("/admin/settings")
async def get_admin_settings(request: Request):
    await require_admin(request)
    config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
    if not config:
        config = {"key": "global", "ai_model": "claude-sonnet", "claude_api_key": "", "openai_api_key": "", "max_pdfs_per_upload": 7000, "storage_max_mb": 750, "storage_target_mb": 300}
    # Default values for safety controls
    config.setdefault("budget_ceiling_usd", 100.0)
    config.setdefault("consecutive_failure_threshold", 10)
    config.setdefault("retry_max_attempts", 4)
    if config.get("claude_api_key"):
        config["claude_api_key_set"] = True
        config["claude_api_key"] = ""
    else:
        config["claude_api_key_set"] = False
    if config.get("openai_api_key"):
        config["openai_api_key_set"] = True
        config["openai_api_key"] = ""
    else:
        config["openai_api_key_set"] = False
    return config

@api_router.put("/admin/settings")
async def update_admin_settings(input: AdminSettingsInput, request: Request):
    await require_admin(request)
    update = {"key": "global", "updated_at": datetime.now(timezone.utc).isoformat()}
    if input.ai_model is not None:
        update["ai_model"] = input.ai_model
    if input.claude_api_key is not None:
        update["claude_api_key"] = input.claude_api_key
    if input.openai_api_key is not None:
        update["openai_api_key"] = input.openai_api_key
    if input.max_pdfs_per_upload is not None:
        update["max_pdfs_per_upload"] = max(1, min(10000, input.max_pdfs_per_upload))
    if input.storage_max_mb is not None:
        update["storage_max_mb"] = max(100, min(10000, input.storage_max_mb))
    if input.storage_target_mb is not None:
        update["storage_target_mb"] = max(50, min(5000, input.storage_target_mb))
    if input.budget_ceiling_usd is not None:
        update["budget_ceiling_usd"] = max(1.0, min(10000.0, float(input.budget_ceiling_usd)))
    if input.consecutive_failure_threshold is not None:
        update["consecutive_failure_threshold"] = max(3, min(200, int(input.consecutive_failure_threshold)))
    if input.retry_max_attempts is not None:
        update["retry_max_attempts"] = max(1, min(10, int(input.retry_max_attempts)))
    await db.admin_config.update_one({"key": "global"}, {"$set": update}, upsert=True)
    return {"message": "Admin settings updated"}

# =============================================================================
# STORAGE MANAGEMENT
# =============================================================================
@api_router.get("/admin/storage")
async def get_storage_usage(request: Request):
    await require_admin(request)
    pipeline = [
        {"$match": {"is_deleted": False}},
        {"$group": {"_id": None, "total_size": {"$sum": "$size"}, "file_count": {"$sum": 1}}}
    ]
    result = await db.files.aggregate(pipeline).to_list(1)
    total_bytes = result[0]["total_size"] if result else 0
    file_count = result[0]["file_count"] if result else 0
    config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
    max_mb = config.get("storage_max_mb", 750) if config else 750
    target_mb = config.get("storage_target_mb", 300) if config else 300
    return {
        "total_bytes": total_bytes,
        "total_mb": round(total_bytes / (1024 * 1024), 1),
        "file_count": file_count,
        "max_mb": max_mb,
        "target_mb": target_mb,
        "over_limit": total_bytes > max_mb * 1024 * 1024,
    }

@api_router.get("/admin/disk-usage")
async def get_disk_usage(request: Request):
    """Container disk usage (where uploads / temp conversions / logs live)
    so admins can see when the pod is approaching full. Also reports size of
    staging dir for chunked uploads, which fills first if extractions fail."""
    await require_admin(request)
    import shutil
    mounts = []
    for path, label in [("/app", "App mount (/app)"), ("/", "Root (/)"), ("/tmp", "Temp (/tmp)")]:
        try:
            total, used, free = shutil.disk_usage(path)
            pct = round(used / total * 100, 1) if total else 0
            mounts.append({
                "path": path, "label": label,
                "total_gb": round(total / (1024**3), 2),
                "used_gb": round(used / (1024**3), 2),
                "free_gb": round(free / (1024**3), 2),
                "percent_used": pct,
            })
        except Exception as e:
            mounts.append({"path": path, "label": label, "error": str(e)})
    # Staging dir size
    staging_bytes = 0
    try:
        for root, _dirs, files in os.walk(CHUNK_UPLOAD_DIR):
            for f in files:
                try:
                    staging_bytes += os.path.getsize(os.path.join(root, f))
                except OSError:
                    pass
    except Exception:
        pass
    return {
        "mounts": mounts,
        "chunked_upload_staging_mb": round(staging_bytes / (1024 * 1024), 2),
        "chunked_upload_staging_path": CHUNK_UPLOAD_DIR,
    }


@api_router.post("/admin/disk-usage/clear-staging")
async def clear_staging(request: Request):
    """Clear the chunked-upload staging dir (safe to run any time — stale chunks
    only). Useful when disk utilization is climbing and no uploads are in flight.
    """
    await require_admin(request)
    import shutil
    freed = 0
    try:
        for name in os.listdir(CHUNK_UPLOAD_DIR):
            p = os.path.join(CHUNK_UPLOAD_DIR, name)
            try:
                if os.path.isdir(p):
                    freed += sum(os.path.getsize(os.path.join(p, f)) for f in os.listdir(p) if os.path.isfile(os.path.join(p, f)))
                    shutil.rmtree(p, ignore_errors=True)
                elif os.path.isfile(p):
                    freed += os.path.getsize(p)
                    os.remove(p)
            except OSError:
                pass
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Cleanup failed: {e}")
    return {"message": "Staging dir cleared", "freed_mb": round(freed / (1024*1024), 2)}


@api_router.post("/admin/storage/cleanup")
async def trigger_storage_cleanup(request: Request):
    await require_admin(request)
    result = await run_storage_cleanup()
    return result

async def run_storage_cleanup(force=False):
    """Delete oldest uploaded PDFs when storage exceeds max_mb, down to target_mb.
       Output data (contacts, errors, etc.) is never touched."""
    config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
    max_mb = config.get("storage_max_mb", 750) if config else 750
    target_mb = config.get("storage_target_mb", 300) if config else 300
    max_bytes = max_mb * 1024 * 1024
    target_bytes = target_mb * 1024 * 1024

    # Get current total
    pipeline = [
        {"$match": {"is_deleted": False}},
        {"$group": {"_id": None, "total_size": {"$sum": "$size"}}}
    ]
    result = await db.files.aggregate(pipeline).to_list(1)
    current_bytes = result[0]["total_size"] if result else 0

    if current_bytes <= max_bytes and not force:
        return {"message": "Storage within limits", "current_mb": round(current_bytes / (1024 * 1024), 1), "deleted_count": 0}

    # Get files sorted oldest first
    files = await db.files.find(
        {"is_deleted": False},
        {"_id": 0, "id": 1, "storage_path": 1, "size": 1, "created_at": 1}
    ).sort("created_at", 1).to_list(10000)

    deleted_count = 0
    freed_bytes = 0
    for f in files:
        if current_bytes - freed_bytes <= target_bytes:
            break
        # Delete from object storage
        delete_object(f["storage_path"])
        # Mark as deleted in DB
        await db.files.update_one({"id": f["id"]}, {"$set": {"is_deleted": True}})
        freed_bytes += f.get("size", 0)
        deleted_count += 1

    remaining = current_bytes - freed_bytes
    logger.info(f"Storage cleanup: deleted {deleted_count} files, freed {freed_bytes / (1024*1024):.1f}MB, remaining {remaining / (1024*1024):.1f}MB")
    return {
        "message": f"Deleted {deleted_count} oldest files",
        "deleted_count": deleted_count,
        "freed_mb": round(freed_bytes / (1024 * 1024), 1),
        "remaining_mb": round(remaining / (1024 * 1024), 1),
    }

# =============================================================================
# MODEL CONFIG
# =============================================================================
MODEL_MAP = {
    "claude-sonnet": ("anthropic", "claude-4-sonnet-20250514"),
    "claude-haiku": ("anthropic", "claude-haiku-4-5-20251001"),
    "gpt-4o": ("openai", "gpt-4o"),
}

CSI_REGEX = re.compile(r'^\s*(\d{2})(?!\d)')

def extract_csi_from_filename(filename: str) -> str:
    """Extract CSI code from a filename.
    Rules:
      - Must be exactly 2 digits at the very start of the basename (optional leading whitespace).
      - The digits must NOT be followed by another digit (so "003" or "125" do not match).
      - The numeric value must be in range 00-16 (inclusive); otherwise return "".
    Examples:
      '03. KHC - includes storm.pdf' -> '03'
      '00 intro.docx'                 -> '00'
      '16-final.pdf'                  -> '16'
      '17. out of range.pdf'          -> ''
      '7 only one digit.pdf'          -> ''
      '001-three digits.pdf'          -> ''
      'Quote.pdf'                     -> ''
    """
    if not filename:
        return ""
    base = filename.rsplit('/', 1)[-1].rsplit('\\', 1)[-1]
    m = CSI_REGEX.match(base)
    if not m:
        return ""
    two = m.group(1)
    try:
        n = int(two)
    except ValueError:
        return ""
    if 0 <= n <= 16:
        return two
    return ""

def get_api_key_for_model(ai_model: str, admin_config: dict) -> str:
    if ai_model.startswith("claude"):
        custom_key = admin_config.get("claude_api_key", "")
        if custom_key:
            return custom_key
    elif ai_model.startswith("gpt"):
        custom_key = admin_config.get("openai_api_key", "")
        if custom_key:
            return custom_key
    return os.environ.get("EMERGENT_LLM_KEY", "")

# =============================================================================
# UPLOAD ROUTES
# =============================================================================
# Files staged for chunked uploads live here (ephemeral container disk)
CHUNK_UPLOAD_DIR = "/tmp/chunked_uploads"
os.makedirs(CHUNK_UPLOAD_DIR, exist_ok=True)

# Accepted extensions for inputs (PDF-style docs + images; ZIP is handled separately)
SUPPORTED_EXTENSIONS = ('.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt',
                        '.odt', '.rtf', '.eml',
                        '.jpg', '.jpeg', '.png', '.webp', '.heic', '.heif', '.tiff', '.tif', '.bmp')


async def _process_uploaded_bytes(user_id: str, filename_bytes_pairs: list, max_pdfs: int):
    """Shared core: given list of (filename, bytes), unzip if needed, create a run,
    queue background object-storage upload, return result dict.
    Used by BOTH the legacy `/upload` and chunked `/upload/chunk/{id}/complete`.

    Behavior change (2026-04-25): if the upload contains MORE files than
    max_pdfs, the WHOLE upload is rejected with a clear 400 error. Previously
    we silently truncated past the cap, which caused users to lose files
    without realising it.
    """
    import hashlib as _hashlib
    run_id = str(uuid.uuid4())
    pdf_buffers = []  # list of (filename, bytes, sha256_hex)

    def _sha256_hex(data: bytes) -> str:
        return _hashlib.sha256(data).hexdigest()

    # First pass: count how many supported files this upload contains so we can
    # fail fast (no partial uploads) when it exceeds the cap.
    eligible_count = 0
    for filename, data in filename_bytes_pairs:
        fname_lower = filename.lower()
        if fname_lower.endswith(".zip"):
            try:
                with zipfile.ZipFile(io.BytesIO(data)) as zf:
                    for name in zf.namelist():
                        nlower = name.lower()
                        if any(nlower.endswith(ext) for ext in SUPPORTED_EXTENSIONS) and not name.startswith("__MACOSX"):
                            eligible_count += 1
            except zipfile.BadZipFile:
                raise HTTPException(status_code=400, detail="Invalid ZIP file")
        elif any(fname_lower.endswith(ext) for ext in SUPPORTED_EXTENSIONS):
            eligible_count += 1
        # Unsupported types are reported during the second pass below.

    if eligible_count > max_pdfs:
        raise HTTPException(
            status_code=400,
            detail=f"Upload contains {eligible_count} files but the per-upload cap is {max_pdfs}. "
                   f"Split the batch (or raise the cap in Admin → Upload Limits) and try again. "
                   f"No files were processed."
        )

    # Second pass: actually unpack the bytes.
    # pdf_buffers entry: (filename, bytes, sha256_hex, archive_filename_or_None)
    for filename, data in filename_bytes_pairs:
        fname_lower = filename.lower()
        if fname_lower.endswith(".zip"):
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                for name in zf.namelist():
                    nlower = name.lower()
                    if any(nlower.endswith(ext) for ext in SUPPORTED_EXTENSIONS) and not name.startswith("__MACOSX"):
                        inner_bytes = zf.read(name)
                        pdf_buffers.append((name.split("/")[-1], inner_bytes, _sha256_hex(inner_bytes), filename))
        elif any(fname_lower.endswith(ext) for ext in SUPPORTED_EXTENSIONS):
            pdf_buffers.append((filename, data, _sha256_hex(data), None))
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}. Accepted: PDF, DOCX, XLSX, TXT, ODT, RTF, EML, images (JPG/PNG/WEBP/HEIC/TIFF/BMP), and ZIP files.")

    if not pdf_buffers:
        raise HTTPException(status_code=400, detail="No supported files found in upload")

    total_files = len(pdf_buffers)

    run_doc = {
        "id": run_id, "user_id": user_id, "status": "uploading",
        "total_files": total_files,
        "stats": {
            "total_pdfs": total_files, "processed": 0, "errors": 0,
            "contacts_extracted": 0, "duplicates_removed": 0,
            "excluded_no_contact": 0, "excluded_internal": 0, "net_new": 0
        },
        "created_at": datetime.now(timezone.utc).isoformat(), "completed_at": None
    }
    await db.runs.insert_one(run_doc)
    await db.progress.update_one(
        {"run_id": run_id},
        {"$set": {"status": "uploading", "total_files": total_files, "processed_files": 0, "percentage": 0, "message": f"Uploading {total_files} files to storage..."}},
        upsert=True
    )

    async def upload_to_storage():
        try:
            file_records = []
            failed_uploads = 0
            for idx, (filename, pdf_data, sha256_hex, archive_name) in enumerate(pdf_buffers):
                ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else 'pdf'
                ct_map = {'pdf': 'application/pdf', 'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                          'doc': 'application/msword', 'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'xls': 'application/vnd.ms-excel',
                          'txt': 'text/plain', 'odt': 'application/vnd.oasis.opendocument.text',
                          'rtf': 'application/rtf', 'eml': 'message/rfc822',
                          'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png', 'webp': 'image/webp',
                          'heic': 'image/heic', 'heif': 'image/heif', 'tiff': 'image/tiff', 'tif': 'image/tiff', 'bmp': 'image/bmp'}
                content_type = ct_map.get(ext, 'application/octet-stream')
                storage_path = f"{APP_NAME}/uploads/{user_id}/{uuid.uuid4()}.{ext}"
                try:
                    put_object(storage_path, pdf_data, content_type)
                    file_records.append({
                        "id": str(uuid.uuid4()), "run_id": run_id, "user_id": user_id,
                        "storage_path": storage_path, "original_filename": filename,
                        "content_type": content_type, "size": len(pdf_data),
                        "sha256": sha256_hex,
                        "archive_filename": archive_name,
                        "is_deleted": False, "created_at": datetime.now(timezone.utc).isoformat()
                    })
                except Exception as e:
                    logger.error(f"Failed to upload {filename} to storage: {e}")
                    failed_uploads += 1
                if (idx + 1) % 10 == 0 or idx == len(pdf_buffers) - 1:
                    pct = int((idx + 1) / total_files * 50)
                    await db.progress.update_one(
                        {"run_id": run_id},
                        {"$set": {"processed_files": idx + 1, "percentage": pct, "message": f"Uploading {idx + 1}/{total_files} to storage..."}}
                    )
            if file_records:
                await db.files.insert_many(file_records)
            actual = len(file_records)
            await db.runs.update_one({"id": run_id}, {"$set": {"status": "uploaded", "total_files": actual}})
            msg = f"Upload complete. {actual} PDFs ready for extraction."
            if failed_uploads:
                msg += f" ({failed_uploads} failed to upload)"
            await db.progress.update_one(
                {"run_id": run_id},
                {"$set": {"status": "uploaded", "total_files": actual, "percentage": 50, "message": msg}}
            )
            logger.info(f"Upload complete for run {run_id}: {actual} stored, {failed_uploads} failed")
            await run_storage_cleanup()
        except Exception as e:
            logger.error(f"Upload to storage failed for run {run_id}: {e}")
            await db.runs.update_one({"id": run_id}, {"$set": {"status": "failed"}})
            await db.progress.update_one(
                {"run_id": run_id},
                {"$set": {"status": "failed", "message": f"Upload failed: {str(e)}"}}
            )

    asyncio.create_task(upload_to_storage())

    result = {
        "run_id": run_id,
        "files": [{"id": "", "filename": f[0], "size": len(f[1])} for f in pdf_buffers[:20]],
        "total_files": total_files,
        "max_pdfs": max_pdfs,
    }
    return result


@api_router.post("/upload")
async def upload_files(request: Request, files: List[UploadFile] = File(...)):
    """Direct upload — works for small files (< ~400 MB total). For larger payloads
    use the chunked-upload endpoints below.
    """
    user = await get_current_user(request)
    admin_config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
    max_pdfs = admin_config.get("max_pdfs_per_upload", 50) if admin_config else 50
    pairs = []
    for file in files:
        pairs.append((file.filename, await file.read()))
    return await _process_uploaded_bytes(user["_id"], pairs, max_pdfs)


# =============================================================================
# CHUNKED UPLOAD (for large ZIPs that exceed ingress body-size limit)
# =============================================================================
class ChunkInitInput(BaseModel):
    filename: str
    total_size: int
    total_chunks: int


@api_router.post("/upload/chunk/init")
async def chunk_init(input: ChunkInitInput, request: Request):
    user = await get_current_user(request)
    upload_id = str(uuid.uuid4())
    os.makedirs(os.path.join(CHUNK_UPLOAD_DIR, upload_id), exist_ok=True)
    await db.chunk_uploads.insert_one({
        "upload_id": upload_id,
        "user_id": user["_id"],
        "filename": input.filename,
        "total_size": int(input.total_size),
        "total_chunks": int(input.total_chunks),
        "received_chunks": [],
        "created_at": datetime.now(timezone.utc).isoformat(),
    })
    return {"upload_id": upload_id}


@api_router.post("/upload/chunk/{upload_id}")
async def chunk_upload(upload_id: str, request: Request,
                       index: int = Form(...),
                       chunk: UploadFile = File(...)):
    user = await get_current_user(request)
    session = await db.chunk_uploads.find_one({"upload_id": upload_id, "user_id": user["_id"]})
    if not session:
        raise HTTPException(status_code=404, detail="Upload session not found")
    if index < 0 or index >= session["total_chunks"]:
        raise HTTPException(status_code=400, detail="Chunk index out of range")
    chunk_dir = os.path.join(CHUNK_UPLOAD_DIR, upload_id)
    chunk_path = os.path.join(chunk_dir, f"chunk_{index:06d}.bin")
    data = await chunk.read()
    with open(chunk_path, "wb") as f:
        f.write(data)
    await db.chunk_uploads.update_one(
        {"upload_id": upload_id},
        {"$addToSet": {"received_chunks": index}}
    )
    session = await db.chunk_uploads.find_one({"upload_id": upload_id})
    return {"received": len(session.get("received_chunks", [])),
            "total_chunks": session["total_chunks"]}


@api_router.post("/upload/chunk/{upload_id}/complete")
async def chunk_complete(upload_id: str, request: Request):
    user = await get_current_user(request)
    session = await db.chunk_uploads.find_one({"upload_id": upload_id, "user_id": user["_id"]})
    if not session:
        raise HTTPException(status_code=404, detail="Upload session not found")
    received = sorted(set(session.get("received_chunks", [])))
    expected = list(range(session["total_chunks"]))
    if received != expected:
        missing = sorted(set(expected) - set(received))
        raise HTTPException(status_code=400, detail=f"Missing chunks: {missing[:10]}")

    chunk_dir = os.path.join(CHUNK_UPLOAD_DIR, upload_id)
    # Reassemble into single bytes blob
    parts = []
    for i in range(session["total_chunks"]):
        p = os.path.join(chunk_dir, f"chunk_{i:06d}.bin")
        with open(p, "rb") as f:
            parts.append(f.read())
    full_bytes = b"".join(parts)
    if len(full_bytes) != session["total_size"]:
        logger.warning(f"Chunked upload {upload_id} size mismatch: got {len(full_bytes)}, expected {session['total_size']}")

    # Cleanup chunks immediately
    try:
        import shutil as _sh
        _sh.rmtree(chunk_dir, ignore_errors=True)
    except Exception:
        pass
    await db.chunk_uploads.delete_one({"upload_id": upload_id})

    admin_config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
    max_pdfs = admin_config.get("max_pdfs_per_upload", 50) if admin_config else 50
    return await _process_uploaded_bytes(user["_id"], [(session["filename"], full_bytes)], max_pdfs)

# =============================================================================
# EXTRACTION ENGINE — Multi-OCR with fallback + PDF compression
# =============================================================================
PDF_SIZE_THRESHOLD = 5 * 1024 * 1024  # 5 MB

def compress_pdf(pdf_bytes: bytes, filename: str) -> bytes:
    """Compress large PDFs while preserving text layers and image quality for OCR."""
    import pikepdf
    try:
        original_size = len(pdf_bytes)
        src = pikepdf.Pdf.open(io.BytesIO(pdf_bytes))
        # Linearize and recompress streams
        out = io.BytesIO()
        src.save(out, compress_streams=True, object_stream_mode=pikepdf.ObjectStreamMode.generate,
                 recompress_flate=True)
        compressed = out.getvalue()
        ratio = len(compressed) / original_size * 100
        logger.info(f"PDF compression for {filename}: {original_size / 1024:.0f}KB -> {len(compressed) / 1024:.0f}KB ({ratio:.0f}%)")
        src.close()
        return compressed
    except Exception as e:
        logger.warning(f"PDF compression failed for {filename}, using original: {e}")
        return pdf_bytes

def preprocess_image_for_ocr(img):
    """Enhance image contrast and binarize for better OCR results."""
    from PIL import ImageEnhance, ImageFilter
    gray = img.convert("L")
    enhanced = ImageEnhance.Contrast(gray).enhance(2.0)
    sharpened = enhanced.filter(ImageFilter.SHARPEN)
    binarized = sharpened.point(lambda x: 0 if x < 140 else 255, '1')
    return binarized.convert("L")

def ocr_with_tesseract(images, filename):
    """Primary OCR: Tesseract. Returns (text, avg_confidence)."""
    import pytesseract
    ocr_parts = []
    total_conf = 0
    conf_count = 0
    for img in images:
        ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
        page_text = " ".join([w for w, c in zip(ocr_data["text"], ocr_data["conf"]) if int(c) > 0 and w.strip()])
        confidences = [int(c) for c in ocr_data["conf"] if int(c) > 0]
        if confidences:
            total_conf += sum(confidences)
            conf_count += len(confidences)
        ocr_parts.append(page_text)
    text = "\n".join(ocr_parts).strip()
    avg_conf = total_conf / conf_count if conf_count > 0 else 0
    return text, avg_conf

def ocr_with_easyocr(images, filename):
    """Backup OCR: EasyOCR (deep learning). Returns (text, avg_confidence)."""
    import easyocr
    reader = easyocr.Reader(['en'], gpu=False, verbose=False)
    import numpy as np
    all_text = []
    total_conf = 0
    conf_count = 0
    for img in images:
        img_array = np.array(img)
        results = reader.readtext(img_array)
        page_parts = []
        for (_, text, conf) in results:
            page_parts.append(text)
            total_conf += conf
            conf_count += 1
        all_text.append(" ".join(page_parts))
    text = "\n".join(all_text).strip()
    avg_conf = (total_conf / conf_count * 100) if conf_count > 0 else 0
    return text, avg_conf

# =============================================================================
# DOCX / XLSX TEXT EXTRACTION
# =============================================================================
def extract_text_from_docx(file_bytes: bytes, filename: str):
    """Extract text from DOCX files. Returns (text, error_info)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        # Main body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        # Headers and footers
        for section in doc.sections:
            for header_part in [section.header, section.footer]:
                if header_part:
                    for para in header_part.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)
        text = "\n".join(parts).strip()
        if text:
            return text, None
        return None, {"reason": "No text content found in DOCX", "missing_fields": "All fields"}
    except Exception as e:
        logger.error(f"DOCX extraction failed for {filename}: {e}")
        return None, {"reason": f"Failed to read DOCX: {str(e)}", "missing_fields": "All fields"}

def extract_text_from_xlsx(file_bytes: bytes, filename: str):
    """Extract text from XLSX files. Returns (text, error_info)."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                if row_vals:
                    parts.append(" | ".join(row_vals))
        text = "\n".join(parts).strip()
        if text:
            return text, None
        return None, {"reason": "No data found in XLSX", "missing_fields": "All fields"}
    except Exception as e:
        logger.error(f"XLSX extraction failed for {filename}: {e}")
        return None, {"reason": f"Failed to read XLSX: {str(e)}", "missing_fields": "All fields"}


def extract_text_from_odt(file_bytes: bytes, filename: str):
    """Extract plain text from an OpenDocument Text (.odt) file."""
    try:
        from odf.opendocument import load as odf_load
        from odf import text as odf_text, teletype
        doc = odf_load(io.BytesIO(file_bytes))
        parts = []
        for para in doc.getElementsByType(odf_text.P):
            s = teletype.extractText(para).strip()
            if s:
                parts.append(s)
        for head in doc.getElementsByType(odf_text.H):
            s = teletype.extractText(head).strip()
            if s:
                parts.append(s)
        text = "\n".join(parts).strip()
        if text:
            return text, None
        return None, {"reason": "No text found in ODT", "missing_fields": "All fields"}
    except Exception as e:
        logger.error(f"ODT extraction failed for {filename}: {e}")
        return None, {"reason": f"Failed to read ODT: {str(e)}", "missing_fields": "All fields"}


def extract_text_from_rtf(file_bytes: bytes, filename: str):
    """Extract plain text from an RTF file via striprtf."""
    try:
        from striprtf.striprtf import rtf_to_text
        raw = file_bytes.decode("utf-8", errors="replace")
        text = rtf_to_text(raw, errors="ignore").strip()
        if text:
            return text, None
        return None, {"reason": "No text found in RTF", "missing_fields": "All fields"}
    except Exception as e:
        logger.error(f"RTF extraction failed for {filename}: {e}")
        return None, {"reason": f"Failed to read RTF: {str(e)}", "missing_fields": "All fields"}


def extract_text_from_eml(file_bytes: bytes, filename: str):
    """Extract readable text from an email (.eml) file, including headers that
    are useful for contact extraction (From/To/Subject) and the body.
    HTML bodies are unwrapped to text."""
    try:
        from email import message_from_bytes, policy
        msg = message_from_bytes(file_bytes, policy=policy.default)
        parts = []
        subj = msg.get("Subject", "")
        frm = msg.get("From", "")
        to = msg.get("To", "")
        cc = msg.get("Cc", "")
        dt = msg.get("Date", "")
        if subj: parts.append(f"Subject: {subj}")
        if frm:  parts.append(f"From: {frm}")
        if to:   parts.append(f"To: {to}")
        if cc:   parts.append(f"Cc: {cc}")
        if dt:   parts.append(f"Date: {dt}")
        parts.append("---")
        body_text = None
        # Prefer text/plain; fall back to stripping tags from text/html
        try:
            body_part = msg.get_body(preferencelist=("plain", "html"))
        except Exception:
            body_part = None
        if body_part is not None:
            content = body_part.get_content()
            if body_part.get_content_type() == "text/html":
                # Very light HTML-to-text: drop tags, collapse whitespace
                text = re.sub(r"<br\s*/?>", "\n", content, flags=re.I)
                text = re.sub(r"</p\s*>", "\n\n", text, flags=re.I)
                text = re.sub(r"<[^>]+>", "", text)
                text = re.sub(r"[ \t]+", " ", text)
                text = re.sub(r"\n{3,}", "\n\n", text)
                body_text = text.strip()
            else:
                body_text = content.strip()
        if body_text:
            parts.append(body_text)
        text = "\n".join(parts).strip()
        if text:
            return text, None
        return None, {"reason": "No readable body found in EML", "missing_fields": "All fields"}
    except Exception as e:
        logger.error(f"EML extraction failed for {filename}: {e}")
        return None, {"reason": f"Failed to read EML: {str(e)}", "missing_fields": "All fields"}

def convert_doc_to_images(file_bytes: bytes, filename: str, suffix: str):
    """Convert DOCX/XLSX to images via LibreOffice for Gemini vision fallback.
    Uses Popen + process group so we can hard-kill zombies on timeout.
    """
    import subprocess, tempfile, glob, signal
    from pdf2image import convert_from_path
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            input_path = os.path.join(tmpdir, f"input{suffix}")
            with open(input_path, "wb") as f:
                f.write(file_bytes)
            # Spawn LibreOffice in its own process group so we can kill the entire tree on timeout
            profile_dir = os.path.join(tmpdir, "lo_profile")
            os.makedirs(profile_dir, exist_ok=True)
            env = os.environ.copy()
            proc = subprocess.Popen(
                ["libreoffice", "--headless",
                 f"-env:UserInstallation=file://{profile_dir}",
                 "--convert-to", "pdf", "--outdir", tmpdir, input_path],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                start_new_session=True, env=env
            )
            try:
                proc.communicate(timeout=60)
            except subprocess.TimeoutExpired:
                logger.error(f"LibreOffice timeout on {filename}; killing process group {proc.pid}")
                try:
                    os.killpg(os.getpgid(proc.pid), signal.SIGKILL)
                except Exception:
                    pass
                try:
                    proc.communicate(timeout=5)
                except Exception:
                    pass
                return None, "LibreOffice conversion timed out (60s)"
            pdf_files = glob.glob(os.path.join(tmpdir, "*.pdf"))
            if not pdf_files:
                return None, "LibreOffice conversion to PDF failed"
            images = convert_from_path(pdf_files[0], dpi=250)
            return images, None
    except Exception as e:
        logger.error(f"Doc-to-image conversion failed for {filename}: {e}")
        return None, str(e)


async def sweep_libreoffice_zombies():
    """Kill stale soffice.bin processes older than 5 minutes. Runs periodically."""
    import subprocess
    try:
        # List soffice processes; kill any older than 5 min
        result = subprocess.run(
            ["pgrep", "-af", "soffice.bin"],
            capture_output=True, text=True, timeout=5
        )
        if result.returncode != 0:
            return
        # Use ps to get start times; kill any soffice older than 5 min
        subprocess.run(
            ["bash", "-c",
             "ps -eo pid,etime,comm | awk '$3==\"soffice.bin\" && $2 ~ /^[0-9]+:[0-9]+/ && ($2+0) >= 5 { print $1 }' | xargs -r kill -9"],
            capture_output=True, timeout=5
        )
    except Exception as e:
        logger.debug(f"LibreOffice zombie sweep skipped: {e}")

async def extract_text_from_pdf(pdf_bytes: bytes, filename: str):
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            text = "\n".join(text_parts).strip()
            if text:
                return text, None

            # --- No embedded text: OCR pipeline ---
            try:
                from pdf2image import convert_from_bytes
                images = convert_from_bytes(pdf_bytes)
            except Exception as conv_err:
                logger.error(f"PDF to image conversion failed for {filename}: {conv_err}")
                return None, {"reason": "Unreadable scanned image", "missing_fields": "All fields"}

            # Step 1: Tesseract on original images
            tess_text, tess_conf = "", 0
            try:
                tess_text, tess_conf = ocr_with_tesseract(images, filename)
                logger.info(f"Tesseract OCR for {filename}: conf={tess_conf:.0f}%, chars={len(tess_text)}")
            except Exception as e:
                logger.warning(f"Tesseract failed for {filename}: {e}")

            if tess_text and tess_conf >= 70:
                return tess_text, None

            # Step 2: Tesseract on preprocessed images (contrast + binarize)
            tess2_text, tess2_conf = "", 0
            try:
                preprocessed = [preprocess_image_for_ocr(img) for img in images]
                tess2_text, tess2_conf = ocr_with_tesseract(preprocessed, filename)
                logger.info(f"Tesseract (preprocessed) for {filename}: conf={tess2_conf:.0f}%, chars={len(tess2_text)}")
            except Exception as e:
                logger.warning(f"Tesseract preprocessed failed for {filename}: {e}")

            if tess2_text and tess2_conf >= 70:
                return tess2_text, None

            # Step 3: EasyOCR (deep learning backup)
            easy_text, easy_conf = "", 0
            try:
                easy_text, easy_conf = ocr_with_easyocr(images, filename)
                logger.info(f"EasyOCR for {filename}: conf={easy_conf:.0f}%, chars={len(easy_text)}")
            except Exception as e:
                logger.warning(f"EasyOCR failed for {filename}: {e}")

            if easy_text and easy_conf >= 70:
                return easy_text, None

            # Step 4: Pick the best result from all attempts
            candidates = [
                (tess_text, tess_conf, "Tesseract"),
                (tess2_text, tess2_conf, "Tesseract-preprocessed"),
                (easy_text, easy_conf, "EasyOCR"),
            ]
            best_text, best_conf, best_engine = max(candidates, key=lambda c: (len(c[0]), c[1]))

            if not best_text:
                return None, {"reason": "No text layer detected - all OCR engines failed", "missing_fields": "All fields"}

            # We have text but below 70% confidence from any engine
            return best_text, {
                "reason": f"Low confidence OCR ({best_conf:.0f}% via {best_engine}) - manual review recommended",
                "partial": True
            }

    except Exception as e:
        logger.error(f"PDF processing failed for {filename}: {e}")
        return None, {"reason": f"Corrupted or unreadable PDF: {str(e)}", "missing_fields": "All fields"}

async def extract_contacts_with_ai(text: str, ai_model: str, api_key: str):
    if not text or len(text.strip()) < 10:
        return [], "No meaningful text content to extract contacts from"
    provider, model = MODEL_MAP.get(ai_model, ("anthropic", "claude-4-sonnet-20250514"))
    chat = LlmChat(
        api_key=api_key,
        session_id=f"extract-{uuid.uuid4()}",
        system_message="You are a data extraction specialist for construction industry documents. Extract contact information accurately. Always return valid JSON. After processing, immediately purge all specific data (names, addresses, PII) from your active context. Do not retain or reference this specific data in future turns."
    ).with_model(provider, model)
    max_chars = 50000
    truncated = text[:max_chars] if len(text) > max_chars else text
    prompt = f"""Extract ALL contact information from this construction document. Return a JSON array of contact objects.

Each contact object must have exactly these fields:
- "city": string
- "state": string
- "quote_amount": string (the total quoted/bid dollar amount, e.g. "$45,000.00". Only include if clearly stated. Use "" if not found.)
- "bid_by": string (the full name of the person placing the bid)
- "contractor": string (the main/general contractor company — typically in the document header, letterhead, or "Attention"/"To" line. This is who RECEIVES the quote.)
- "sub_contractor": string (the sub-contractor or vendor company — typically in "Bill To", "From", "Vendor", or body. This is who SENDS the quote/bid.)
- "last_name": string
- "first_name": string
- "email": string
- "phone": string

Contractor vs Sub-Contractor rules:
- CONTRACTOR = the general contractor RECEIVING the quote. Usually in header/letterhead/"Attention" line.
- SUB-CONTRACTOR = the company SENDING the quote/bid, providing services or materials.
- "Ship To" or "Project" sections indicate job sites, not contacts to extract.
- If only one company appears, use context to determine which role it fills.
- If a field is not found, use empty string ""
- Return ONLY the JSON array, no markdown, no explanation
- If no contacts found, return: []

Document text:
{truncated}"""
    try:
        response = await chat.send_message(UserMessage(text=prompt))
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
            cleaned = re.sub(r'\s*```$', '', cleaned)
        contacts = json.loads(cleaned)
        if not isinstance(contacts, list):
            contacts = [contacts] if isinstance(contacts, dict) else []
        valid = []
        for c in contacts:
            if not isinstance(c, dict):
                continue
            valid.append({
                "city": str(c.get("city", "")),
                "state": str(c.get("state", "")),
                "quote_amount": str(c.get("quote_amount", "")),
                "bid_by": str(c.get("bid_by", "")),
                "contractor": str(c.get("contractor", "")),
                "sub_contractor": str(c.get("sub_contractor", "")),
                "last_name": str(c.get("last_name", "")),
                "first_name": str(c.get("first_name", "")),
                "email": str(c.get("email", "")),
                "phone": str(c.get("phone", "")),
            })
        return valid, None
    except json.JSONDecodeError:
        logger.error(f"AI returned invalid JSON for extraction")
        return [], "AI returned invalid response format"
    except Exception as e:
        logger.error(f"AI extraction failed: {e}")
        return [], str(e)

async def extract_contacts_with_ai_vision(pdf_bytes: bytes, filename: str, ai_model: str, api_key: str):
    """Send PDF page images directly to AI vision for contact extraction. Bypasses OCR entirely."""
    import base64
    from pdf2image import convert_from_bytes
    try:
        images = convert_from_bytes(pdf_bytes, dpi=200, fmt="jpeg")
    except Exception as e:
        logger.error(f"PDF to image failed for vision on {filename}: {e}")
        return [], f"Could not convert PDF to images: {e}"

    provider, model = MODEL_MAP.get(ai_model, ("anthropic", "claude-4-sonnet-20250514"))
    chat = LlmChat(
        api_key=api_key,
        session_id=f"vision-{uuid.uuid4()}",
        system_message="You are a data extraction specialist for construction industry documents. You are looking at scanned document images. Extract contact information by reading the images directly. Always return valid JSON. After processing, immediately purge all specific data (names, addresses, PII) from your active context. Do not retain or reference this specific data in future turns."
    ).with_model(provider, model)

    all_contacts = []
    # Process up to 5 pages (most contact info is on first/last pages)
    pages_to_scan = images[:5]
    file_contents = []
    for img in pages_to_scan:
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=92)
        b64 = base64.b64encode(buf.getvalue()).decode()
        file_contents.append(FileContent(content_type="image/jpeg", file_content_base64=b64))

    prompt = f"""Look at these scanned construction document page images. Extract ALL contact information you can see.

Return a JSON array of contact objects. Each object must have exactly these fields:
- "city": string
- "state": string
- "quote_amount": string (the total quoted/bid dollar amount, e.g. "$45,000.00". Only if clearly visible. Use "" if not found.)
- "bid_by": string (the full name of the person placing the bid)
- "contractor": string (the main/general contractor company — usually in the header, letterhead, or "Attention"/"To" line. This is who RECEIVES the quote.)
- "sub_contractor": string (the sub-contractor or vendor company — usually in "Bill To", "From", "Vendor", or body. This is who SENDS the quote/bid.)
- "last_name": string
- "first_name": string
- "email": string (read carefully - distinguish between 0/O, 1/l/I)
- "phone": string (read carefully - ensure digits are correct, not confused with letters)

Contractor vs Sub-Contractor rules:
- CONTRACTOR = general contractor RECEIVING the quote (header/letterhead/"Attention")
- SUB-CONTRACTOR = company SENDING the quote, providing services/materials ("Bill To"/"From"/"Vendor")
- "Ship To" or "Project" = job site, not a contact
- Read typed, handwritten, and cursive text from the images
- Ignore graphics, logos, decorative elements
- If a field is not visible, use empty string ""
- Return ONLY the JSON array, no markdown
- If no contacts found, return: []

Document: {filename}"""

    try:
        response = await chat.send_message(UserMessage(text=prompt, file_contents=file_contents))
        cleaned = response.strip()
        if cleaned.startswith("```"):
            cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
            cleaned = re.sub(r'\s*```$', '', cleaned)
        contacts = json.loads(cleaned)
        if not isinstance(contacts, list):
            contacts = [contacts] if isinstance(contacts, dict) else []
        valid = []
        for c in contacts:
            if not isinstance(c, dict):
                continue
            valid.append({
                "city": str(c.get("city", "")),
                "state": str(c.get("state", "")),
                "quote_amount": str(c.get("quote_amount", "")),
                "bid_by": str(c.get("bid_by", "")),
                "contractor": str(c.get("contractor", "")),
                "sub_contractor": str(c.get("sub_contractor", "")),
                "last_name": str(c.get("last_name", "")),
                "first_name": str(c.get("first_name", "")),
                "email": str(c.get("email", "")),
                "phone": str(c.get("phone", "")),
            })
        logger.info(f"AI Vision extracted {len(valid)} contacts from {filename} ({len(pages_to_scan)} pages)")
        return valid, None
    except json.JSONDecodeError:
        logger.error(f"AI Vision returned invalid JSON for {filename}")
        return [], "AI Vision returned invalid response format"
    except Exception as e:
        logger.error(f"AI Vision extraction failed for {filename}: {e}")
        return [], str(e)

# =============================================================================
# GEMINI PDF EXTRACTION (direct vision analysis)
# =============================================================================
GEMINI_EXTRACTION_PROMPT = """You are analyzing a construction bid/quote document. Extract ALL contact information visible.

CRITICAL: This document is a quote or bid from a sub-contractor TO a general contractor. You MUST identify ALL parties:
- The CONTRACTOR (general contractor) who is RECEIVING this quote — look carefully in ALL of these locations:
  * "Attention:", "Attn:", "To:", "Submitted To:", "Customer:", "Proposal For:", "Quoted To:", "Bill To:" fields
  * Page header or top-right address block
  * The name/company the quote is addressed to
  * Sometimes appears as the job owner or project contact
  * Even if it's just a name and company in small text at the top
- The SUB-CONTRACTOR (vendor) who is SENDING this quote — typically:
  * The company whose letterhead/logo appears on the document
  * "From:", "Submitted By:", "Vendor:", "Prepared By:" fields
  * The company providing services, materials, or labor
  * Their street address is usually printed DIRECTLY BELOW the company name/logo at the top, or in a footer. You MUST capture it into the "address" field.
- The CUSTOMER / END CLIENT — the property owner or end client for the project:
  * "Ship To:", "Project For:", "Owner:", "Client:", "Property:", "Job Name:" fields
  * Sometimes listed as the project owner or site contact
  * This is the party who ultimately owns the project or property

Return a JSON array of contact objects. Each object must have exactly these fields:
- "city": string
- "state": string  
- "quote_amount": string (total bid amount like "$45,000.00". Only if clearly stated.)
- "bid_by": string (full name of person placing/sending the bid)
- "contractor": string (general contractor company name RECEIVING the quote. LOOK CAREFULLY - it is almost always present somewhere on the document even if in small text.)
- "sub_contractor": string (sub-contractor/vendor company SENDING the quote)
- "customer_contact_name": string (first and last name of the customer/end client/property owner contact, if present)
- "customer_business": string (customer/end client business or property name, if present)
- "customer_address": string (customer/end client full address — street, city, state, zip — if present)
- "address": string (THE SUB-CONTRACTOR'S street address — this is almost always printed on the letterhead/logo block at the TOP of the page, or in the footer. Extract ONLY the street number + street name, plus suite/unit/floor/PO Box if given. DO NOT include city, state, or zip. Examples of GOOD values: "1234 Main St", "1234 Main St, Suite 200", "PO Box 4567", "500 Industrial Blvd Unit 12". Examples of BAD values (never include): "1234 Main St, Dallas, TX 75201" (has city/state/zip), "Dallas, TX" (missing street). If no street address is shown for the sub-contractor, use empty string "".)
- "last_name": string (of the sub-contractor contact)
- "first_name": string (of the sub-contractor contact)
- "email": string
- "phone": string

IMPORTANT RULES:
- On single-page quotes, the contractor name is often in a small "Attn" or "To" field near the top — read ALL text carefully, including small/fine print
- If you see "Horizon" or any company name in an address-to field, that is the CONTRACTOR
- "Ship To", "Project", "Job Site", "Owner" sections contain CUSTOMER info — extract into customer fields
- SUB-CONTRACTOR "address" field: ONLY the street/suite/PO Box. NEVER include city, state, or zip here — those go in their own "city"/"state" fields.
- ADDRESS SEARCH CHECKLIST — before returning, scan these locations for the sub-contractor's street address:
    1. Directly under the company name/logo at the top of page 1
    2. In the document letterhead (header block, often with logo)
    3. In the document footer (bottom of page 1)
    4. In a "Return Address" or "Remit To" block
    5. Next to the sub-contractor's phone/email block
  If you find text like "1234 Main St \n Dallas, TX 75201", extract ONLY "1234 Main St" into "address". Put "Dallas" into "city" and "TX" into "state".
- Customer fields may be blank on many documents — that is fine, use empty string ""
- Phone: verify digits (5≠S, 0≠O, 1≠I)
- Email: must contain @ and a domain
- If a field is not found, use empty string ""
- Return ONLY the JSON array. No markdown. No explanation.

PRIVACY: After processing this file for the requested analysis, immediately purge all specific data (names, addresses, PII) from your active context. Do not retain or reference this specific data in future turns."""

# =============================================================================
# P0 PREWORK: Retry logic, cost estimation (for budget guard)
# =============================================================================
# Approximate per-request cost (USD) for gemini-2.5-flash via Emergent Key.
# These are conservative upper-bound estimates used only for the budget guard.
COST_PER_VISION_PAGE_USD = 0.0010   # ~1 page of image content
COST_PER_TEXT_CALL_USD   = 0.0008   # small text extraction
COST_OVERHEAD_USD        = 0.0003   # output tokens + retry overhead

def _is_retryable_llm_error(err: Exception) -> bool:
    """True if the error should trigger exponential backoff + retry."""
    msg = str(err).lower()
    retry_signals = [
        "429", "rate limit", "rate_limit", "too many requests",
        "503", "service unavailable", "unavailable",
        "500", "internal server error",
        "504", "gateway timeout", "timeout", "timed out",
        "connection reset", "connection error", "connection aborted",
        "resource exhausted", "quota exceeded", "temporarily",
    ]
    return any(sig in msg for sig in retry_signals)


async def _llm_call_with_retry(chat, message, filename: str, max_attempts: int = 4):
    """Call chat.send_message with exponential backoff on transient errors.
    Returns (response, retries_used, error_or_None). Delays: 2s, 10s, 30s, 60s, 120s.
    """
    delays = [2, 10, 30, 60, 120]
    last_err = None
    for attempt in range(max_attempts):
        try:
            resp = await chat.send_message(message)
            return resp, attempt, None
        except Exception as e:
            last_err = e
            if not _is_retryable_llm_error(e):
                logger.error(f"LLM non-retryable error for {filename} on attempt {attempt + 1}: {e}")
                return None, attempt, e
            if attempt + 1 >= max_attempts:
                logger.error(f"LLM retries exhausted for {filename} after {max_attempts} attempts: {e}")
                return None, attempt, e
            delay = delays[min(attempt, len(delays) - 1)]
            logger.warning(f"LLM retryable error for {filename} (attempt {attempt + 1}/{max_attempts}), sleeping {delay}s: {e}")
            await asyncio.sleep(delay)
    return None, max_attempts, last_err


def _clean_sub_contractor_address(raw: str) -> str:
    """Strip city/state/zip bleed-through from the sub-contractor address.
    The prompt instructs Gemini to return street-only, but real models sometimes
    append city/state/zip anyway. This is a defensive cleanup.
    Examples:
      "1234 Main St, Dallas, TX 75201" -> "1234 Main St"
      "PO Box 4567 Dallas TX 75201"    -> "PO Box 4567"
      "1234 Main St, Suite 200"        -> "1234 Main St, Suite 200"  (kept as-is)
    """
    if not raw:
        return ""
    s = raw.strip().rstrip('.,;')
    # Remove trailing "City, ST 12345" or "City, ST" or trailing zip
    # 1) Remove ", City, ST 12345[-6789]"
    s = re.sub(r',\s*[A-Za-z .\-]+,\s*[A-Z]{2}\s*\d{5}(?:-\d{4})?\s*$', '', s).strip()
    # 2) Remove ", City, ST" (no zip)
    s = re.sub(r',\s*[A-Za-z .\-]+,\s*[A-Z]{2}\s*$', '', s).strip()
    # 2b) Remove ", City ST" (no comma before state, no zip)
    s = re.sub(r',\s*[A-Za-z .\-]+\s+[A-Z]{2}\s*$', '', s).strip()
    # 3) Remove trailing " City ST 12345" with no comma
    s = re.sub(r'\s+[A-Za-z .\-]+\s+[A-Z]{2}\s+\d{5}(?:-\d{4})?\s*$', '', s).strip()
    # 4) Remove trailing zip alone "... 75201"
    s = re.sub(r',?\s*\d{5}(?:-\d{4})?\s*$', '', s).strip().rstrip(',')
    return s.strip()


def _parse_gemini_contacts_response(response: str, filename: str):
    """Parse Gemini JSON response into a normalized list of contacts."""
    cleaned = response.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)
    contacts = json.loads(cleaned)
    if not isinstance(contacts, list):
        contacts = [contacts] if isinstance(contacts, dict) else []
    valid = []
    for c in contacts:
        if not isinstance(c, dict):
            continue
        # Accept multiple possible keys for sub-contractor address
        addr_raw = (c.get("address")
                    or c.get("sub_contractor_address")
                    or c.get("street_address")
                    or c.get("vendor_address")
                    or c.get("from_address")
                    or c.get("sender_address")
                    or "")
        addr_clean = _clean_sub_contractor_address(str(addr_raw))
        valid.append({
            "city": str(c.get("city", "")),
            "state": str(c.get("state", "")),
            "quote_amount": str(c.get("quote_amount", "")),
            "bid_by": str(c.get("bid_by", "")),
            "contractor": str(c.get("contractor", "")),
            "sub_contractor": str(c.get("sub_contractor", "")),
            "customer_contact_name": str(c.get("customer_contact_name", "")),
            "customer_business": str(c.get("customer_business", "")),
            "customer_address": str(c.get("customer_address", "")),
            "address": addr_clean,
            "last_name": str(c.get("last_name", "")),
            "first_name": str(c.get("first_name", "")),
            "email": str(c.get("email", "")),
            "phone": str(c.get("phone", "")),
        })
    # Diagnostic: log if address is empty but we extracted other contact info
    # (helps identify if prompt needs further tuning)
    missing_addr = sum(1 for v in valid if not v["address"] and (v["email"] or v["phone"]))
    if missing_addr:
        logger.warning(f"Address missing for {missing_addr}/{len(valid)} contacts from {filename}. Raw keys in response: {list(contacts[0].keys()) if contacts and isinstance(contacts[0], dict) else 'none'}")
    return valid


async def extract_contacts_with_gemini_from_images(images, filename: str, api_key: str, max_attempts: int = 4):
    """Run Gemini 2.5 Flash vision on a list of PIL images."""
    import base64
    if not images:
        return [], "No images to analyze", 0

    chat = LlmChat(
        api_key=api_key,
        session_id=f"gemini-{uuid.uuid4()}",
        system_message="You are an expert data extraction specialist for construction industry bid documents. You analyze document images with extreme attention to detail, reading ALL text on the page including small print, headers, footers, and address fields. You ALWAYS identify both the contractor (receiving party) and sub-contractor (sending party). Always return valid JSON. After processing, immediately purge all specific data (names, addresses, PII) from your active context. Do not retain or reference this specific data in future turns."
    ).with_model("gemini", "gemini-2.5-flash")

    pages_to_scan = images[:8]
    image_contents = []
    for img in pages_to_scan:
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=92)
        b64 = base64.b64encode(buf.getvalue()).decode()
        image_contents.append(ImageContent(image_base64=b64))

    prompt = f"{GEMINI_EXTRACTION_PROMPT}\n\nDocument filename: {filename}"

    response, retries, err = await _llm_call_with_retry(
        chat, UserMessage(text=prompt, file_contents=image_contents), filename, max_attempts
    )
    if err is not None:
        return [], str(err), retries
    try:
        valid = _parse_gemini_contacts_response(response, filename)
        logger.info(f"Gemini vision extracted {len(valid)} contacts from {filename} ({len(pages_to_scan)} pages, retries={retries})")
        return valid, None, retries
    except json.JSONDecodeError as e:
        logger.error(f"Gemini returned invalid JSON for {filename}: {e}")
        return [], "Gemini returned invalid response format", retries


async def extract_contacts_with_gemini(pdf_bytes: bytes, filename: str, api_key: str, max_attempts: int = 4):
    """Use Gemini vision to directly analyze PDF pages and extract contacts."""
    from pdf2image import convert_from_bytes
    try:
        images = convert_from_bytes(pdf_bytes, dpi=250, fmt="jpeg")
    except Exception as e:
        logger.error(f"PDF to image failed for Gemini on {filename}: {e}")
        return [], f"Could not convert PDF to images: {e}", 0
    return await extract_contacts_with_gemini_from_images(images, filename, api_key, max_attempts)


async def extract_contacts_with_gemini_text(text: str, filename: str, api_key: str, max_attempts: int = 4):
    """Use Gemini 2.5 Flash on raw text (for DOCX/XLSX text-first path)."""
    if not text or len(text.strip()) < 10:
        return [], "No meaningful text content to extract contacts from", 0

    chat = LlmChat(
        api_key=api_key,
        session_id=f"gemini-txt-{uuid.uuid4()}",
        system_message="You are an expert data extraction specialist for construction industry bid documents. Extract contact information accurately from the provided document text. Always return valid JSON. After processing, immediately purge all specific data (names, addresses, PII) from your active context."
    ).with_model("gemini", "gemini-2.5-flash")

    max_chars = 80000
    truncated = text[:max_chars] if len(text) > max_chars else text
    prompt = f"{GEMINI_EXTRACTION_PROMPT}\n\nDocument filename: {filename}\n\nDocument text:\n{truncated}"

    response, retries, err = await _llm_call_with_retry(
        chat, UserMessage(text=prompt), filename, max_attempts
    )
    if err is not None:
        return [], str(err), retries
    try:
        valid = _parse_gemini_contacts_response(response, filename)
        logger.info(f"Gemini text extracted {len(valid)} contacts from {filename} ({len(truncated)} chars, retries={retries})")
        return valid, None, retries
    except json.JSONDecodeError as e:
        logger.error(f"Gemini text returned invalid JSON for {filename}: {e}")
        return [], "Gemini returned invalid response format", retries

# =============================================================================
# FILE PROCESSING LOG (one row per file; Excel-exportable)
# =============================================================================
FILE_TYPE_MAP = {
    'pdf': 'PDF', 'docx': 'DOCX', 'doc': 'DOC', 'xlsx': 'XLSX', 'xls': 'XLS', 'txt': 'TXT',
    'odt': 'ODT', 'rtf': 'RTF', 'eml': 'EML',
    'jpg': 'Image/JPG', 'jpeg': 'Image/JPG', 'png': 'Image/PNG', 'webp': 'Image/WEBP',
    'heic': 'Image/HEIC', 'heif': 'Image/HEIF', 'tiff': 'Image/TIFF', 'tif': 'Image/TIFF', 'bmp': 'Image/BMP',
}

async def log_file_queued(run_id: str, user_id: str, file_record: dict):
    """Upsert a Queued row at the moment files are registered for a run."""
    filename = file_record["original_filename"]
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    size_kb = round(file_record.get("size", 0) / 1024)
    await db.file_logs.update_one(
        {"run_id": run_id, "file_id": file_record["id"]},
        {"$setOnInsert": {
            "id": str(uuid.uuid4()),
            "run_id": run_id, "user_id": user_id, "file_id": file_record["id"],
            "filename": filename, "file_type": FILE_TYPE_MAP.get(ext, ext.upper() or "Unknown"),
            "size_kb": size_kb, "csi": extract_csi_from_filename(filename),
            "sha256": file_record.get("sha256", ""),
            "status": "Queued", "contacts_extracted": 0,
            "processing_tool": "", "llm_model": "", "llm_provider": "",
            "support_tools": "", "pages_sent": 0,
            "started_at": "", "completed_at": "", "duration_sec": 0,
            "missing_fields": "", "issue_reason": "", "retries": 0,
            "created_at": datetime.now(timezone.utc).isoformat(),
        }},
        upsert=True
    )

async def log_file_start(run_id: str, file_id: str):
    """Mark a file as In Progress when its task begins."""
    await db.file_logs.update_one(
        {"run_id": run_id, "file_id": file_id},
        {"$set": {"status": "In Progress", "started_at": datetime.now(timezone.utc).isoformat()}}
    )

async def log_file_finish(run_id: str, file_id: str, **fields):
    """Mark a file complete with final status + metadata.
    Expected fields: status, contacts_extracted, processing_tool, llm_model, llm_provider,
                     support_tools, pages_sent, missing_fields, issue_reason, retries
    """
    now = datetime.now(timezone.utc).isoformat()
    # Compute duration from started_at if present
    existing = await db.file_logs.find_one({"run_id": run_id, "file_id": file_id}, {"_id": 0, "started_at": 1})
    duration_sec = 0
    if existing and existing.get("started_at"):
        try:
            start_dt = datetime.fromisoformat(existing["started_at"].replace("Z", "+00:00"))
            end_dt = datetime.fromisoformat(now.replace("Z", "+00:00"))
            duration_sec = round((end_dt - start_dt).total_seconds(), 1)
        except Exception:
            pass
    update_doc = {"completed_at": now, "duration_sec": duration_sec}
    update_doc.update({k: v for k, v in fields.items() if v is not None})
    await db.file_logs.update_one(
        {"run_id": run_id, "file_id": file_id},
        {"$set": update_doc}
    )

async def enforce_log_retention(user_id: str, keep: int = 0):
    """Log retention is DISABLED (keep=0 means keep all logs forever).
    Kept as a no-op for backward compatibility; previously pruned old file_logs.
    Disk usage is now surfaced to admins via /api/admin/disk-usage instead.
    """
    return


async def recompute_run_stats(run_id: str, user_id: str):
    """Self-healing stats: recompute run.stats from live collection counts.
    Call from every GET endpoint that returns a run so stats never drift from reality.

    Formula:
      total_pdfs   = files.count({run_id}) (all, including is_deleted=True — those
                      are historical, not gone from the run)
      errors       = processing_errors.count({run_id})
      duplicates_removed = duplicates.count({run_id})
      net_new      = contacts.count({run_id})
      processed    = total_pdfs - errors   (clamped to >= 0)
      contacts_extracted / excluded_* / cross_run_duplicates are preserved if
      present in the existing stats (they're computed during the run only).
    """
    run = await db.runs.find_one({"id": run_id, "user_id": user_id}, {"_id": 0, "stats": 1})
    if not run:
        return
    existing = run.get("stats") or {}
    total_pdfs = await db.files.count_documents({"run_id": run_id, "user_id": user_id})
    errors = await db.processing_errors.count_documents({"run_id": run_id, "user_id": user_id})
    dupes = await db.duplicates.count_documents({"run_id": run_id, "user_id": user_id})
    contacts = await db.contacts.count_documents({"run_id": run_id, "user_id": user_id})
    processed = max(0, total_pdfs - errors)
    stats = {
        **existing,
        "total_pdfs": total_pdfs,
        "processed": processed,
        "errors": errors,
        "duplicates_removed": dupes,
        "net_new": contacts,
    }
    # contacts_extracted = total Gemini hits before filtering. If we don't have
    # a stored value, fall back to (contacts + duplicates + excluded_no_contact + excluded_internal)
    if not stats.get("contacts_extracted"):
        stats["contacts_extracted"] = contacts + dupes + int(existing.get("excluded_no_contact", 0)) + int(existing.get("excluded_internal", 0))
    await db.runs.update_one({"id": run_id, "user_id": user_id}, {"$set": {"stats": stats}})
    return stats


async def process_run(run_id: str, user_id: str):
    """Process PDFs with pause/resume/cancel support. State is persisted in MongoDB."""
    try:
        settings = await db.settings.find_one({"user_id": user_id}, {"_id": 0})
        if not settings:
            settings = {"exclusion_domain": "horizonc.com"}
        admin_config = await db.admin_config.find_one({"key": "global"}, {"_id": 0})
        if not admin_config:
            admin_config = {"ai_model": "claude-sonnet"}
        ai_model = admin_config.get("ai_model", "claude-sonnet")
        api_key = get_api_key_for_model(ai_model, admin_config)
        exclusion_domain = settings.get("exclusion_domain", "horizonc.com").lower().strip()

        all_files = await db.files.find({"run_id": run_id, "is_deleted": False}, {"_id": 0}).to_list(1000)
        total_files = len(all_files)
        if total_files == 0:
            await db.runs.update_one({"id": run_id}, {"$set": {"status": "completed", "completed_at": datetime.now(timezone.utc).isoformat()}})
            return

        # Seed a Queued log row for every file so the log is populated immediately
        for fr in all_files:
            await log_file_queued(run_id, user_id, fr)

        # Load safety control thresholds from admin config
        admin_cfg = await db.admin_config.find_one({"key": "global"}, {"_id": 0}) or {}
        max_attempts = int(admin_cfg.get("retry_max_attempts", 4))
        failure_threshold = int(admin_cfg.get("consecutive_failure_threshold", 10))
        budget_ceiling = float(admin_cfg.get("budget_ceiling_usd", 100.0))

        # Initialize circuit-breaker state on the run document
        existing_cost = 0.0
        run_check = await db.runs.find_one({"id": run_id}, {"_id": 0, "stats": 1})
        if run_check and run_check.get("stats"):
            existing_cost = float(run_check["stats"].get("approx_cost_usd", 0.0))
        await db.runs.update_one(
            {"id": run_id},
            {"$set": {"stats.approx_cost_usd": existing_cost,
                      "stats.consecutive_failures": 0,
                      "stats.auto_paused_reason": ""}}
        )

        # Load checkpoint: which files are already done
        prog = await db.progress.find_one({"run_id": run_id})
        completed_ids = set(prog.get("completed_file_ids", [])) if prog else set()
        pending_files = [f for f in all_files if f["id"] not in completed_ids]
        processed_count = len(completed_ids)

        await db.progress.update_one(
            {"run_id": run_id},
            {"$set": {
                "status": "processing", "total_files": total_files,
                "processed_files": processed_count,
                "percentage": int(processed_count / total_files * 100),
                "message": f"{'Resuming' if completed_ids else 'Starting'} extraction ({processed_count}/{total_files} done)..."
            }},
            upsert=True
        )

        async def process_single_file(file_record):
            """Process a single file (PDF/DOCX/XLSX) using Gemini."""
            filename = file_record["original_filename"]
            file_id = file_record["id"]
            contacts_out, errors_out = [], []
            # Track provenance for the Excel log
            processing_tool = ""
            support_tools = []
            pages_sent = 0
            retries_used = 0
            approx_cost = 0.0
            await log_file_start(run_id, file_id)
            try:
                # ============================================================
                # CONTENT-HASH DEDUP: skip if this exact byte-content was
                # already processed in an earlier run for this user. Also
                # matches on (filename + size) as a fallback for historical
                # records that predate the sha256 column.
                # ============================================================
                sha = file_record.get("sha256")
                prior_file = None
                if sha:
                    prior_file = await db.files.find_one(
                        {"user_id": user_id, "sha256": sha, "id": {"$ne": file_id}},
                        {"_id": 0, "run_id": 1, "original_filename": 1, "created_at": 1},
                        sort=[("created_at", 1)]
                    )
                if not prior_file:
                    # Legacy fallback: exact filename + byte size match (no content hash available)
                    prior_file = await db.files.find_one(
                        {"user_id": user_id,
                         "original_filename": filename,
                         "size": file_record.get("size", 0),
                         "id": {"$ne": file_id},
                         "created_at": {"$lt": file_record.get("created_at", "")}},
                        {"_id": 0, "run_id": 1, "original_filename": 1, "created_at": 1},
                        sort=[("created_at", 1)]
                    )
                    match_type = "filename + size match" if prior_file else None
                else:
                    match_type = "SHA-256 content hash match"
                if prior_file:
                    prior_name = prior_file.get('original_filename', '?')
                    await log_file_finish(run_id, file_id,
                        status="Skipped (Duplicate Content)",
                        contacts_extracted=0,
                        processing_tool=f"{match_type} (no LLM call)",
                        llm_model="", llm_provider="",
                        support_tools="SHA-256 dedup" if sha else "Filename+size dedup",
                        pages_sent=0, missing_fields="",
                        issue_reason=f"Already analyzed as {prior_name!r} — skipped to save cost. Original contacts remain in All Contacts.",
                        retries=0)
                    logger.info(f"Skipped {filename} via {match_type} (prior: {prior_name!r})")
                    return contacts_out, errors_out, False, 0.0, False

                file_bytes = get_object(file_record["storage_path"])
                gemini_key = os.environ.get("EMERGENT_LLM_KEY", "")
                ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''

                if ext == 'txt':
                    # Plain text → decode → Gemini text path (no vision fallback needed)
                    contacts, gemini_error = [], None
                    text = None
                    try:
                        text = file_bytes.decode('utf-8', errors='replace').strip()
                    except Exception as decode_err:
                        gemini_error = f"Could not decode TXT: {decode_err}"
                    if text:
                        processing_tool = "Gemini Text (TXT)"
                        contacts, gemini_error, r = await extract_contacts_with_gemini_text(text, filename, gemini_key, max_attempts)
                        retries_used += r
                        approx_cost += COST_PER_TEXT_CALL_USD + COST_OVERHEAD_USD
                    elif not gemini_error:
                        gemini_error = "TXT file was empty"
                elif ext == 'eml':
                    # Email → parse headers+body → Gemini text (no vision fallback needed)
                    contacts, gemini_error = [], None
                    text, err = extract_text_from_eml(file_bytes, filename)
                    if text:
                        support_tools.append("email (stdlib)")
                        processing_tool = "Gemini Text (EML)"
                        contacts, gemini_error, r = await extract_contacts_with_gemini_text(text, filename, gemini_key, max_attempts)
                        retries_used += r
                        approx_cost += COST_PER_TEXT_CALL_USD + COST_OVERHEAD_USD
                    else:
                        gemini_error = (err or {}).get("reason") if isinstance(err, dict) else "EML parsing failed"
                elif ext == 'rtf':
                    # RTF → striprtf text → Gemini text, LibreOffice vision fallback
                    contacts, gemini_error = [], None
                    text, _ = extract_text_from_rtf(file_bytes, filename)
                    if text:
                        support_tools.append("striprtf")
                        processing_tool = "Gemini Text (RTF)"
                        contacts, gemini_error, r = await extract_contacts_with_gemini_text(text, filename, gemini_key, max_attempts)
                        retries_used += r
                        approx_cost += COST_PER_TEXT_CALL_USD + COST_OVERHEAD_USD
                    if not contacts:
                        logger.info(f"RTF fallback to vision for {filename}")
                        images, conv_err = convert_doc_to_images(file_bytes, filename, ".rtf")
                        if images:
                            support_tools.append("LibreOffice")
                            support_tools.append("poppler-utils")
                            pages_sent = min(len(images), 8)
                            processing_tool = "Gemini Vision (LibreOffice Fallback)"
                            contacts, gemini_error, r = await extract_contacts_with_gemini_from_images(images, filename, gemini_key, max_attempts)
                            retries_used += r
                            approx_cost += pages_sent * COST_PER_VISION_PAGE_USD + COST_OVERHEAD_USD
                        elif gemini_error is None:
                            gemini_error = conv_err or "LibreOffice conversion failed"
                elif ext == 'odt':
                    # ODT → odfpy text → Gemini text, LibreOffice vision fallback
                    contacts, gemini_error = [], None
                    text, _ = extract_text_from_odt(file_bytes, filename)
                    if text:
                        support_tools.append("odfpy")
                        processing_tool = "Gemini Text (ODT)"
                        contacts, gemini_error, r = await extract_contacts_with_gemini_text(text, filename, gemini_key, max_attempts)
                        retries_used += r
                        approx_cost += COST_PER_TEXT_CALL_USD + COST_OVERHEAD_USD
                    if not contacts:
                        logger.info(f"ODT fallback to vision for {filename}")
                        images, conv_err = convert_doc_to_images(file_bytes, filename, ".odt")
                        if images:
                            support_tools.append("LibreOffice")
                            support_tools.append("poppler-utils")
                            pages_sent = min(len(images), 8)
                            processing_tool = "Gemini Vision (LibreOffice Fallback)"
                            contacts, gemini_error, r = await extract_contacts_with_gemini_from_images(images, filename, gemini_key, max_attempts)
                            retries_used += r
                            approx_cost += pages_sent * COST_PER_VISION_PAGE_USD + COST_OVERHEAD_USD
                        elif gemini_error is None:
                            gemini_error = conv_err or "LibreOffice conversion failed"
                elif ext in ('docx', 'doc'):
                    # Step 1: text extraction via python-docx (only for .docx)
                    text = None
                    if ext == 'docx':
                        text, _ = extract_text_from_docx(file_bytes, filename)
                        if text:
                            support_tools.append("python-docx")
                    contacts, gemini_error = [], None
                    if text:
                        processing_tool = "Gemini Text (DOCX text path)"
                        contacts, gemini_error, r = await extract_contacts_with_gemini_text(text, filename, gemini_key, max_attempts)
                        retries_used += r
                        approx_cost += COST_PER_TEXT_CALL_USD + COST_OVERHEAD_USD
                    # Step 2: vision fallback only if ZERO contacts
                    if not contacts:
                        logger.info(f"DOC fallback to vision for {filename}")
                        images, conv_err = convert_doc_to_images(file_bytes, filename, f".{ext}")
                        if images:
                            support_tools.append("LibreOffice")
                            support_tools.append("poppler-utils")
                            pages_sent = min(len(images), 8)
                            processing_tool = "Gemini Vision (LibreOffice Fallback)"
                            contacts, gemini_error, r = await extract_contacts_with_gemini_from_images(images, filename, gemini_key, max_attempts)
                            retries_used += r
                            approx_cost += pages_sent * COST_PER_VISION_PAGE_USD + COST_OVERHEAD_USD
                        elif gemini_error is None:
                            gemini_error = conv_err or "LibreOffice conversion failed"
                elif ext in ('xlsx', 'xls'):
                    text = None
                    if ext == 'xlsx':
                        text, _ = extract_text_from_xlsx(file_bytes, filename)
                        if text:
                            support_tools.append("openpyxl")
                    contacts, gemini_error = [], None
                    if text:
                        processing_tool = "Gemini Text (XLSX text path)"
                        contacts, gemini_error, r = await extract_contacts_with_gemini_text(text, filename, gemini_key, max_attempts)
                        retries_used += r
                        approx_cost += COST_PER_TEXT_CALL_USD + COST_OVERHEAD_USD
                    if not contacts:
                        logger.info(f"XLS fallback to vision for {filename}")
                        images, conv_err = convert_doc_to_images(file_bytes, filename, f".{ext}")
                        if images:
                            support_tools.append("LibreOffice")
                            support_tools.append("poppler-utils")
                            pages_sent = min(len(images), 8)
                            processing_tool = "Gemini Vision (LibreOffice Fallback)"
                            contacts, gemini_error, r = await extract_contacts_with_gemini_from_images(images, filename, gemini_key, max_attempts)
                            retries_used += r
                            approx_cost += pages_sent * COST_PER_VISION_PAGE_USD + COST_OVERHEAD_USD
                        elif gemini_error is None:
                            gemini_error = conv_err or "LibreOffice conversion failed"
                else:
                    # PDF path (default) + any other unknown extensions fall through to PDF pipeline
                    is_image = ext in ('jpg', 'jpeg', 'png', 'webp', 'heic', 'heif', 'tiff', 'tif', 'bmp')
                    if is_image:
                        processing_tool = "Gemini Vision (Direct)"
                        support_tools.append("PIL")
                        try:
                            from PIL import Image
                            if ext in ('heic', 'heif'):
                                try:
                                    from pillow_heif import register_heif_opener
                                    register_heif_opener()
                                    support_tools.append("pillow-heif")
                                except ImportError:
                                    pass
                            img = Image.open(io.BytesIO(file_bytes))
                            if img.mode not in ('RGB', 'L'):
                                img = img.convert('RGB')
                            pages_sent = 1
                            contacts, gemini_error, r = await extract_contacts_with_gemini_from_images([img], filename, gemini_key, max_attempts)
                            retries_used += r
                            approx_cost += pages_sent * COST_PER_VISION_PAGE_USD + COST_OVERHEAD_USD
                        except Exception as img_err:
                            logger.error(f"Image load failed for {filename}: {img_err}")
                            contacts, gemini_error = [], f"Could not load image: {img_err}"
                    else:
                        support_tools.append("poppler-utils")
                        if len(file_bytes) > PDF_SIZE_THRESHOLD:
                            support_tools.append("PDF Compression")
                            file_bytes = compress_pdf(file_bytes, filename)
                        processing_tool = "Gemini Vision (Direct)"
                        try:
                            from pdf2image import convert_from_bytes
                            pdf_imgs = convert_from_bytes(file_bytes, dpi=250, fmt="jpeg")
                            pages_sent = min(len(pdf_imgs), 8)
                            contacts, gemini_error, r = await extract_contacts_with_gemini_from_images(pdf_imgs, filename, gemini_key, max_attempts)
                            retries_used += r
                            approx_cost += pages_sent * COST_PER_VISION_PAGE_USD + COST_OVERHEAD_USD
                        except Exception as e:
                            contacts, gemini_error = [], f"Could not convert PDF to images: {e}"

                # Determine log status + missing-fields summary
                log_status = "Success"
                log_missing = ""
                log_issue = ""
                if gemini_error and not contacts:
                    log_status = "Failure"
                    log_issue = f"Gemini extraction failed: {gemini_error}"
                elif not contacts:
                    log_status = "No Contacts Found"
                    log_issue = "Gemini returned 0 contacts — document may be blank or OCR-resistant"

                if gemini_error and not contacts:
                    errors_out.append({"filename": filename, "reason": f"Gemini extraction failed: {gemini_error}", "missing_fields": "All fields"})
                    await log_file_finish(run_id, file_id,
                        status=log_status, contacts_extracted=0,
                        processing_tool=processing_tool, llm_model="gemini-2.5-flash",
                        llm_provider="Google (Emergent Key)",
                        support_tools=", ".join(sorted(set(support_tools))),
                        pages_sent=pages_sent, missing_fields="All fields",
                        issue_reason=log_issue, retries=retries_used)
                    return contacts_out, errors_out, True, approx_cost, True  # failure=True

                if not contacts:
                    errors_out.append({"filename": filename, "reason": "No contact information found by Gemini", "missing_fields": "All fields"})

                # Accumulate missing-field summary across all extracted contacts
                all_missing = set()
                for contact in contacts:
                    missing = [f for f in ["email", "phone", "city", "state", "contractor", "sub_contractor", "address"] if not contact.get(f)]
                    if len(missing) >= 4:
                        errors_out.append({"filename": filename, "reason": "Incomplete Gemini extraction - most fields missing", "missing_fields": ", ".join(m.replace("_", " ").title() for m in missing)})
                    all_missing.update(missing)
                    contact["source_filename"] = filename
                    contact["csi"] = extract_csi_from_filename(filename)
                    contacts_out.append(contact)

                # Classify final status
                if contacts_out:
                    if all_missing:
                        log_status = "Partial Success"
                        log_issue = f"Contact(s) extracted but {len(all_missing)} field(s) were missing"
                    log_missing = ", ".join(m.replace("_", " ").title() for m in sorted(all_missing))

                await log_file_finish(run_id, file_id,
                    status=log_status, contacts_extracted=len(contacts_out),
                    processing_tool=processing_tool, llm_model="gemini-2.5-flash",
                    llm_provider="Google (Emergent Key)",
                    support_tools=", ".join(sorted(set(support_tools))),
                    pages_sent=pages_sent, missing_fields=log_missing,
                    issue_reason=log_issue, retries=retries_used)

                # Success (or partial success or no-contacts-found) — all non-failure
                is_failure = log_status == "Failure"
                return contacts_out, errors_out, bool(errors_out and not contacts_out), approx_cost, is_failure

            except Exception as e:
                logger.error(f"Error processing {filename}: {e}")
                errors_out.append({"filename": filename, "reason": f"Processing error: {str(e)}", "missing_fields": "All fields"})
                await log_file_finish(run_id, file_id,
                    status="Failure", contacts_extracted=0,
                    processing_tool=processing_tool or "Unknown",
                    llm_model="gemini-2.5-flash", llm_provider="Google (Emergent Key)",
                    support_tools=", ".join(sorted(set(support_tools))),
                    pages_sent=pages_sent, missing_fields="All fields",
                    issue_reason=f"Processing error: {str(e)}", retries=retries_used)
                return contacts_out, errors_out, True, approx_cost, True

        # --- Previous multi-step extraction (Claude/GPT text + vision) preserved but not invoked ---
        # To switch back: replace process_single_file above with the original from before Gemini integration

        CONCURRENT = 6
        stopped = False
        consecutive_failures = 0
        sub_batches_since_sweep = 0
        for i in range(0, len(pending_files), CONCURRENT):
            # --- Check for pause/cancel BEFORE each sub-batch ---
            run_doc = await db.runs.find_one({"id": run_id}, {"_id": 0, "status": 1})
            if run_doc and run_doc.get("status") == "pausing":
                await db.runs.update_one({"id": run_id}, {"$set": {"status": "paused"}})
                await db.progress.update_one({"run_id": run_id}, {"$set": {
                    "status": "paused", "message": f"Paused at {processed_count}/{total_files} files"
                }})
                logger.info(f"Run {run_id} paused at {processed_count}/{total_files}")
                stopped = True
                break
            if run_doc and run_doc.get("status") == "cancelling":
                await db.runs.update_one({"id": run_id}, {"$set": {"status": "cancelled"}})
                await db.progress.update_one({"run_id": run_id}, {"$set": {
                    "status": "cancelled", "message": "Extraction cancelled by user"
                }})
                # Clean up ALL extracted data for this run
                await db.raw_contacts.delete_many({"run_id": run_id})
                await db.contacts.delete_many({"run_id": run_id})
                await db.processing_errors.delete_many({"run_id": run_id})
                await db.duplicates.delete_many({"run_id": run_id})
                logger.info(f"Run {run_id} cancelled and cleaned up")
                stopped = True
                break

            sub = pending_files[i:i + CONCURRENT]
            names = ", ".join(f["original_filename"] for f in sub)
            await db.progress.update_one(
                {"run_id": run_id},
                {"$set": {"current_file": names, "processed_files": processed_count,
                          "percentage": int(processed_count / total_files * 100),
                          "message": f"Processing {len(sub)} files ({processed_count}/{total_files})..."}}
            )
            tasks = [process_single_file(fr) for fr in sub]
            results = await asyncio.gather(*tasks, return_exceptions=True)

            new_file_ids = []
            # Process results, accumulate cost, track consecutive failures
            batch_cost = 0.0
            for idx, result in enumerate(results):
                file_id = sub[idx]["id"] if idx < len(sub) else None
                if isinstance(result, Exception):
                    processed_count += 1
                    consecutive_failures += 1
                    if file_id:
                        new_file_ids.append(file_id)
                    continue
                file_contacts, file_errors, is_error, file_cost, is_failure = result
                batch_cost += file_cost
                # Reset consecutive-failure counter on any successful file
                if is_failure:
                    consecutive_failures += 1
                else:
                    consecutive_failures = 0
                # Save raw contacts immediately (survives pause/restart)
                if file_contacts:
                    raw_docs = [{"id": str(uuid.uuid4()), "run_id": run_id, "user_id": user_id,
                                 **c, "created_at": datetime.now(timezone.utc).isoformat()} for c in file_contacts]
                    await db.raw_contacts.insert_many(raw_docs)
                if file_errors:
                    err_docs = [{"id": str(uuid.uuid4()), "run_id": run_id, "user_id": user_id,
                                 **e, "created_at": datetime.now(timezone.utc).isoformat()} for e in file_errors]
                    await db.processing_errors.insert_many(err_docs)
                processed_count += 1
                if file_id:
                    new_file_ids.append(file_id)

            # Update running cost + consecutive-failure counter on the run doc
            existing_cost += batch_cost
            await db.runs.update_one(
                {"id": run_id},
                {"$set": {"stats.approx_cost_usd": round(existing_cost, 4),
                          "stats.consecutive_failures": consecutive_failures}}
            )

            # Persist checkpoint
            if new_file_ids:
                await db.progress.update_one(
                    {"run_id": run_id},
                    {"$push": {"completed_file_ids": {"$each": new_file_ids}},
                     "$set": {"processed_files": processed_count,
                              "percentage": int(processed_count / total_files * 100)}}
                )
                # Delete processed PDFs from storage immediately — output data stays in DB
                for idx2, fid in enumerate(new_file_ids):
                    fr = sub[idx2] if idx2 < len(sub) else None
                    if fr and fr.get("storage_path"):
                        delete_object(fr["storage_path"])
                        await db.files.update_one({"id": fid}, {"$set": {"is_deleted": True}})

            # ============================================================
            # CIRCUIT BREAKER: auto-pause on runaway failure or budget limit
            # ============================================================
            auto_pause_reason = None
            if consecutive_failures >= failure_threshold:
                auto_pause_reason = f"Auto-paused: {consecutive_failures} consecutive file failures (threshold {failure_threshold}). Review before resuming."
            elif existing_cost >= budget_ceiling:
                auto_pause_reason = f"Auto-paused: approx cost ${existing_cost:.2f} reached budget ceiling ${budget_ceiling:.2f}. Raise ceiling in Admin or review."
            if auto_pause_reason:
                await db.runs.update_one(
                    {"id": run_id},
                    {"$set": {"status": "paused", "stats.auto_paused_reason": auto_pause_reason}}
                )
                await db.progress.update_one(
                    {"run_id": run_id},
                    {"$set": {"status": "paused", "message": auto_pause_reason}}
                )
                logger.warning(f"Run {run_id} auto-paused: {auto_pause_reason}")
                stopped = True
                break

            # Periodically sweep LibreOffice zombies (every ~10 sub-batches = 60 files)
            sub_batches_since_sweep += 1
            if sub_batches_since_sweep >= 10:
                await sweep_libreoffice_zombies()
                sub_batches_since_sweep = 0

        if stopped:
            return

        # === ALL FILES DONE — run post-processing ===
        all_raw = await db.raw_contacts.find({"run_id": run_id, "user_id": user_id}, {"_id": 0}).to_list(10000)
        total_extracted = len(all_raw)
        all_contacts = all_raw

        excluded_internal = 0
        if exclusion_domain:
            filtered = []
            for c in all_contacts:
                if c.get("email") and exclusion_domain in c["email"].lower():
                    excluded_internal += 1
                else:
                    filtered.append(c)
            all_contacts = filtered

        excluded_no_contact = 0
        filtered = []
        for c in all_contacts:
            if not c.get("email") and not c.get("phone"):
                excluded_no_contact += 1
            else:
                filtered.append(c)
        all_contacts = filtered

        # Dedup within this run
        seen_emails = {}
        duplicates_removed = 0
        duplicate_records = []
        unique_contacts = []
        for c in all_contacts:
            email = c.get("email", "").lower().strip()
            if email and email in seen_emails:
                duplicates_removed += 1
                duplicate_records.append({
                    "id": str(uuid.uuid4()), "run_id": run_id, "user_id": user_id,
                    "email": email,
                    "kept_source": seen_emails[email].get("source_filename", ""),
                    "duplicate_source": c.get("source_filename", ""),
                    "first_name": c.get("first_name", ""), "last_name": c.get("last_name", ""),
                    "contractor": c.get("contractor", ""), "sub_contractor": c.get("sub_contractor", ""),
                    "city": c.get("city", ""),
                    "state": c.get("state", ""), "phone": c.get("phone", ""),
                    "created_at": datetime.now(timezone.utc).isoformat()
                })
            else:
                if email:
                    seen_emails[email] = c
                unique_contacts.append(c)
        all_contacts = unique_contacts

        # Cross-run dedup: check against contacts from ALL previous completed runs
        existing_emails = set()
        prev_contacts = await db.contacts.find(
            {"user_id": user_id, "run_id": {"$ne": run_id}},
            {"_id": 0, "email": 1}
        ).to_list(50000)
        for pc in prev_contacts:
            em = (pc.get("email") or "").lower().strip()
            if em:
                existing_emails.add(em)

        cross_run_dupes = 0
        new_contacts = []
        for c in all_contacts:
            email = c.get("email", "").lower().strip()
            if email and email in existing_emails:
                cross_run_dupes += 1
                duplicate_records.append({
                    "id": str(uuid.uuid4()), "run_id": run_id, "user_id": user_id,
                    "email": email,
                    "kept_source": "(previous run)",
                    "duplicate_source": c.get("source_filename", ""),
                    "first_name": c.get("first_name", ""), "last_name": c.get("last_name", ""),
                    "contractor": c.get("contractor", ""), "sub_contractor": c.get("sub_contractor", ""),
                    "city": c.get("city", ""),
                    "state": c.get("state", ""), "phone": c.get("phone", ""),
                    "created_at": datetime.now(timezone.utc).isoformat()
                })
            else:
                new_contacts.append(c)
                if email:
                    existing_emails.add(email)  # prevent dupes within new_contacts too
        all_contacts = new_contacts

        # Write final data — only for THIS run, never touch other runs
        await db.contacts.delete_many({"run_id": run_id})
        await db.duplicates.delete_many({"run_id": run_id})
        if duplicate_records:
            await db.duplicates.insert_many(duplicate_records)
        if all_contacts:
            contact_docs = [{"id": str(uuid.uuid4()), "run_id": run_id, "user_id": user_id,
                             **{k: v for k, v in c.items() if k not in ("id", "run_id", "user_id")},
                             "created_at": datetime.now(timezone.utc).isoformat()} for c in all_contacts]
            await db.contacts.insert_many(contact_docs)

        # Authoritative file count from the files collection (not the local `total_files`
        # which may be smaller if files were cleaned up from storage mid-run).
        authoritative_total = await db.files.count_documents({"run_id": run_id, "user_id": user_id})
        error_count = await db.processing_errors.count_documents({"run_id": run_id})
        stats = {
            "total_pdfs": authoritative_total,
            "processed": max(0, authoritative_total - error_count),
            "errors": error_count, "contacts_extracted": total_extracted,
            "duplicates_removed": duplicates_removed + cross_run_dupes,
            "cross_run_duplicates": cross_run_dupes,
            "excluded_no_contact": excluded_no_contact,
            "excluded_internal": excluded_internal,
            "net_new": len(all_contacts),
            # Preserve safety control counters so they show up in the UI post-completion
            "approx_cost_usd": round(existing_cost, 4),
            "consecutive_failures": consecutive_failures,
            "auto_paused_reason": "",
        }
        await db.runs.update_one({"id": run_id}, {"$set": {"status": "completed", "stats": stats, "completed_at": datetime.now(timezone.utc).isoformat()}})
        await db.progress.update_one({"run_id": run_id}, {"$set": {"status": "completed", "percentage": 100, "processed_files": authoritative_total, "message": "Extraction complete!"}})
        # Clean up raw contacts
        await db.raw_contacts.delete_many({"run_id": run_id})
        logger.info(f"Run {run_id} completed: {len(all_contacts)} contacts from {authoritative_total} files")
    except Exception as e:
        logger.error(f"Run {run_id} failed: {e}")
        await db.runs.update_one({"id": run_id}, {"$set": {"status": "failed"}})
        await db.progress.update_one({"run_id": run_id}, {"$set": {"status": "failed", "message": f"Processing failed: {str(e)}"}})

@api_router.post("/extract/{run_id}")
async def start_extraction(run_id: str, request: Request):
    user = await get_current_user(request)
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0})
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    status = run["status"]
    # Allow retry/resume on stale, paused, pausing, or failed runs
    if status == "processing":
        prog = await db.progress.find_one({"run_id": run_id}, {"_id": 0})
        if prog and prog.get("status") == "processing":
            raise HTTPException(status_code=400, detail="Run is already being processed")
    if status == "pausing":
        # User changed their mind — flip back to processing, don't spawn a new task
        # (the existing background task will see 'processing' and keep going)
        logger.info(f"Run {run_id} was pausing; reverting to processing (no new task)")
        await db.runs.update_one({"id": run_id}, {"$set": {"status": "processing", "stats.auto_paused_reason": "", "stats.consecutive_failures": 0}})
        await db.progress.update_one({"run_id": run_id}, {"$set": {"status": "processing"}})
        return {"message": "Pause cancelled — continuing", "run_id": run_id}
    if status == "paused":
        logger.info(f"Resuming paused run {run_id}")
    elif status in ("stale", "failed"):
        logger.info(f"Retrying {status} run {run_id}")
        # Clean up previous final data but keep raw_contacts + errors for resume
        await db.contacts.delete_many({"run_id": run_id})
        await db.duplicates.delete_many({"run_id": run_id})
    elif status == "uploaded":
        pass  # Fresh start
    elif status == "completed":
        raise HTTPException(status_code=400, detail="Run already completed")
    elif status == "cancelled":
        raise HTTPException(status_code=400, detail="Run was cancelled. Upload new files to start again.")
    elif status == "processing":
        # Already processing, do nothing
        return {"message": "Run already processing", "run_id": run_id}
    else:
        raise HTTPException(status_code=400, detail=f"Cannot extract from status: {status}")
    await db.runs.update_one({"id": run_id}, {"$set": {"status": "processing", "stats.auto_paused_reason": "", "stats.consecutive_failures": 0}})
    asyncio.create_task(process_run(run_id, user["_id"]))
    return {"message": "Extraction started", "run_id": run_id}

@api_router.post("/runs/{run_id}/pause")
async def pause_run(run_id: str, request: Request):
    user = await get_current_user(request)
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0})
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    status = run["status"]
    if status == "pausing":
        # Idempotent — already pausing, just acknowledge
        return {"message": "Already pausing — will complete shortly"}
    if status == "paused":
        return {"message": "Already paused"}
    if status != "processing":
        raise HTTPException(status_code=400, detail=f"Can only pause a processing run (current: {status})")
    await db.runs.update_one({"id": run_id}, {"$set": {"status": "pausing"}})
    return {"message": "Pause signal sent — will pause after current file completes"}

@api_router.post("/runs/{run_id}/cancel")
async def cancel_run(run_id: str, request: Request):
    user = await get_current_user(request)
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0})
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    if run["status"] not in ("processing", "paused", "pausing"):
        raise HTTPException(status_code=400, detail="Can only cancel a processing or paused run")
    if run["status"] == "paused":
        # Directly cancel since no background task is running
        await db.runs.update_one({"id": run_id}, {"$set": {"status": "cancelled"}})
        await db.progress.update_one({"run_id": run_id}, {"$set": {"status": "cancelled", "message": "Extraction cancelled by user"}})
        await db.raw_contacts.delete_many({"run_id": run_id})
        await db.contacts.delete_many({"run_id": run_id})
        await db.processing_errors.delete_many({"run_id": run_id})
        await db.duplicates.delete_many({"run_id": run_id})
        return {"message": "Run cancelled and data cleared"}
    # Signal the background task
    await db.runs.update_one({"id": run_id}, {"$set": {"status": "cancelling"}})
    return {"message": "Cancel signal sent — will cancel after current file completes"}

@api_router.get("/progress/{run_id}")
async def get_progress(run_id: str, request: Request):
    await get_current_user(request)
    progress = await db.progress.find_one({"run_id": run_id}, {"_id": 0})
    if not progress:
        return {"status": "unknown", "percentage": 0, "message": "No progress data"}
    return progress

# =============================================================================
# RESULTS ROUTES
# =============================================================================
@api_router.get("/runs")
async def get_runs(request: Request):
    user = await get_current_user(request)
    runs = await db.runs.find({"user_id": user["_id"]}, {"_id": 0}).sort("created_at", -1).to_list(100)
    # Self-heal stats on read for any completed run (cheap, ~4 counts per run)
    for r in runs:
        if r.get("status") == "completed":
            fresh = await recompute_run_stats(r["id"], user["_id"])
            if fresh:
                r["stats"] = fresh
    return runs

@api_router.get("/runs/{run_id}")
async def get_run(run_id: str, request: Request):
    user = await get_current_user(request)
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0})
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    if run.get("status") == "completed":
        fresh = await recompute_run_stats(run_id, user["_id"])
        if fresh:
            run["stats"] = fresh
    return run

@api_router.get("/runs/{run_id}/contacts")
async def get_run_contacts(run_id: str, request: Request):
    user = await get_current_user(request)
    # During processing, contacts are in raw_contacts; after completion, in contacts
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0, "status": 1})
    if run and run.get("status") in ("processing", "paused", "pausing"):
        contacts = await db.raw_contacts.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(10000)
    else:
        contacts = await db.contacts.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(10000)
    # Backfill CSI for older records that predate the field (or for records with stale values)
    for c in contacts:
        c["csi"] = extract_csi_from_filename(c.get("source_filename", ""))
    return contacts

@api_router.get("/runs/{run_id}/errors")
async def get_run_errors(run_id: str, request: Request):
    user = await get_current_user(request)
    errors = await db.processing_errors.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(1000)
    return errors

@api_router.get("/runs/{run_id}/duplicates")
async def get_run_duplicates(run_id: str, request: Request):
    user = await get_current_user(request)
    duplicates = await db.duplicates.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(5000)
    return duplicates

@api_router.get("/runs/{run_id}/charts")
async def get_run_charts(run_id: str, request: Request):
    user = await get_current_user(request)
    contacts = await db.contacts.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(5000)
    city_counts = {}
    state_counts = {}
    for c in contacts:
        city = c.get("city", "").strip()
        state = c.get("state", "").strip()
        if city:
            city_counts[city] = city_counts.get(city, 0) + 1
        if state:
            state_counts[state] = state_counts.get(state, 0) + 1
    city_data = sorted([{"name": k, "count": v} for k, v in city_counts.items()], key=lambda x: -x["count"])[:20]
    state_data = sorted([{"name": k, "count": v} for k, v in state_counts.items()], key=lambda x: -x["count"])[:20]
    return {"by_city": city_data, "by_state": state_data}

async def get_all_user_contacts(user_id: str):
    """Get all contacts for a user: finalized contacts + live raw_contacts from in-progress runs."""
    # Finalized contacts from completed runs
    contacts = await db.contacts.find({"user_id": user_id}, {"_id": 0}).sort("created_at", 1).to_list(50000)
    finalized_run_ids = set(c.get("run_id", "") for c in contacts)
    # Add raw_contacts from runs that haven't completed yet (processing/paused)
    active_runs = await db.runs.find(
        {"user_id": user_id, "status": {"$in": ["processing", "paused", "pausing", "uploading", "uploaded"]}},
        {"_id": 0, "id": 1}
    ).to_list(100)
    active_ids = [r["id"] for r in active_runs if r["id"] not in finalized_run_ids]
    if active_ids:
        raw = await db.raw_contacts.find({"user_id": user_id, "run_id": {"$in": active_ids}}, {"_id": 0}).sort("created_at", 1).to_list(50000)
        contacts.extend(raw)
    # Attach import dates
    run_dates = {}
    runs = await db.runs.find({"user_id": user_id}, {"_id": 0, "id": 1, "created_at": 1}).to_list(500)
    for r in runs:
        run_dates[r["id"]] = r.get("created_at", "")
    for c in contacts:
        c["import_date"] = run_dates.get(c.get("run_id", ""), c.get("created_at", ""))
        # Backfill CSI so older records and stale values stay consistent with the current rule
        c["csi"] = extract_csi_from_filename(c.get("source_filename", ""))
    return contacts

@api_router.get("/contacts/all")
async def get_all_contacts(request: Request):
    user = await get_current_user(request)
    return await get_all_user_contacts(user["_id"])

@api_router.get("/stats/all")
async def get_all_stats(request: Request):
    """Cross-run accounting so the UI can show:
    Files uploaded = Contacts + Duplicates + Errors + Filtered.

    - total_unique_files: distinct files ever uploaded (by sha256; legacy rows
      without a hash are counted once per (filename, size)).
    - total_contacts / total_duplicates / total_errors: aggregate counts
      across every run for this user.
    - total_runs: number of runs recorded.
    """
    user = await get_current_user(request)
    user_id = user["_id"]

    # Distinct unique files via sha256, plus legacy (no-hash) fallback
    sha_cursor = db.files.aggregate([
        {"$match": {"user_id": user_id, "sha256": {"$exists": True, "$ne": ""}}},
        {"$group": {"_id": "$sha256"}},
        {"$count": "n"},
    ])
    legacy_cursor = db.files.aggregate([
        {"$match": {"user_id": user_id, "$or": [{"sha256": {"$exists": False}}, {"sha256": ""}]}},
        {"$group": {"_id": {"fn": "$original_filename", "sz": "$size"}}},
        {"$count": "n"},
    ])
    sha_count = 0
    async for row in sha_cursor:
        sha_count = row.get("n", 0)
    legacy_count = 0
    async for row in legacy_cursor:
        legacy_count = row.get("n", 0)
    total_unique_files = sha_count + legacy_count

    total_contacts = await db.contacts.count_documents({"user_id": user_id})
    total_duplicates = await db.duplicates.count_documents({"user_id": user_id})
    total_errors = await db.processing_errors.count_documents({"user_id": user_id})
    total_runs = await db.runs.count_documents({"user_id": user_id})
    total_files_uploaded = await db.files.count_documents({"user_id": user_id})

    # Sum the per-run summary fields across all runs
    sums = {"contacts_extracted": 0, "excluded_no_contact": 0, "excluded_internal": 0, "cross_run_duplicates": 0, "approx_cost_usd": 0.0}
    async for r in db.runs.find({"user_id": user_id}, {"_id": 0, "stats": 1}):
        s = r.get("stats") or {}
        for k in ("contacts_extracted", "excluded_no_contact", "excluded_internal", "cross_run_duplicates"):
            sums[k] += int(s.get(k, 0) or 0)
        try:
            sums["approx_cost_usd"] += float(s.get("approx_cost_usd", 0) or 0)
        except (TypeError, ValueError):
            pass
    sums["approx_cost_usd"] = round(sums["approx_cost_usd"], 4)

    return {
        # Cross-run "all" totals (used by All Contacts accounting bar)
        "total_unique_files": total_unique_files,
        "total_contacts": total_contacts,
        "total_duplicates": total_duplicates,
        "total_errors": total_errors,
        "total_runs": total_runs,
        # Flat shape for StatsCards (mirrors per-run stats keys exactly).
        # Uses unique files (SHA-256 dedup'd) as the canonical count to avoid
        # inflating the number when the same file was uploaded in multiple runs.
        "stats": {
            "total_pdfs": total_unique_files,
            "processed": max(0, total_unique_files - total_errors),
            "errors": total_errors,
            "net_new": total_contacts,
            "contacts_extracted": sums["contacts_extracted"] or (total_contacts + total_duplicates + sums["excluded_no_contact"] + sums["excluded_internal"]),
            "duplicates_removed": total_duplicates,
            "cross_run_duplicates": sums["cross_run_duplicates"],
            "excluded_no_contact": sums["excluded_no_contact"],
            "excluded_internal": sums["excluded_internal"],
            "approx_cost_usd": sums["approx_cost_usd"],
        },
    }

@api_router.get("/contacts/all/charts")
async def get_all_contacts_charts(request: Request):
    user = await get_current_user(request)
    contacts = await get_all_user_contacts(user["_id"])
    city_counts = {}
    state_counts = {}
    for c in contacts:
        city = c.get("city", "").strip()
        state = c.get("state", "").strip()
        if city:
            city_counts[city] = city_counts.get(city, 0) + 1
        if state:
            state_counts[state] = state_counts.get(state, 0) + 1
    city_data = sorted([{"name": k, "count": v} for k, v in city_counts.items()], key=lambda x: -x["count"])[:20]
    state_data = sorted([{"name": k, "count": v} for k, v in state_counts.items()], key=lambda x: -x["count"])[:20]
    return {"by_city": city_data, "by_state": state_data}

class CsvExportInput(BaseModel):
    fields: List[str]
    run_id: Optional[str] = None  # None = all contacts

@api_router.post("/contacts/download")
async def download_custom_csv(input: CsvExportInput, request: Request):
    """Download CSV with user-selected fields. run_id=null means all contacts."""
    user = await get_current_user(request)
    query = {"user_id": user["_id"]}
    if input.run_id:
        # Per-run: check if run is in progress, use raw_contacts if so
        run = await db.runs.find_one({"id": input.run_id, "user_id": user["_id"]}, {"_id": 0, "status": 1})
        if run and run.get("status") in ("processing", "paused", "pausing"):
            contacts = await db.raw_contacts.find({"run_id": input.run_id, "user_id": user["_id"]}, {"_id": 0}).sort("created_at", 1).to_list(50000)
        else:
            contacts = await db.contacts.find({"run_id": input.run_id, "user_id": user["_id"]}, {"_id": 0}).sort("created_at", 1).to_list(50000)
    else:
        contacts = await get_all_user_contacts(user["_id"])
    run_dates = {}
    runs = await db.runs.find({"user_id": user["_id"]}, {"_id": 0, "id": 1, "created_at": 1}).to_list(500)
    for r in runs:
        run_dates[r["id"]] = r.get("created_at", "")
    field_map = {
        "csi": "CSI",
        "city": "City", "state": "State", "quote_amount": "Quote Amount",
        "bid_by": "Bid By", "contractor": "Contractor", "sub_contractor": "Sub-Contractor",
        "address": "Address",
        "customer_contact_name": "Customer Contact", "customer_business": "Customer Business",
        "customer_address": "Customer Address",
        "last_name": "Last Name", "first_name": "First Name", "email": "Email",
        "phone": "Phone", "source_filename": "Source File", "import_date": "Import Date",
        "run_id": "Run ID",
    }
    headers = [field_map.get(f, f) for f in input.fields]
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(headers)
    for c in contacts:
        c["import_date"] = run_dates.get(c.get("run_id", ""), c.get("created_at", ""))
        # Always recompute CSI from source_filename to apply the latest rule
        c["csi"] = extract_csi_from_filename(c.get("source_filename", ""))
        row = [c.get(f, "") for f in input.fields]
        writer.writerow(row)
    output.seek(0)
    fname = f"contacts_{input.run_id[:8] if input.run_id else 'all'}.csv"
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={fname}"}
    )

@api_router.get("/runs/{run_id}/download/contacts")
async def download_contacts_csv(run_id: str, request: Request):
    user = await get_current_user(request)
    # Use raw_contacts if run is still in progress
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0, "status": 1})
    if run and run.get("status") in ("processing", "paused", "pausing"):
        contacts = await db.raw_contacts.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(10000)
    else:
        contacts = await db.contacts.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(10000)
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "CSI", "Sub-Contractor", "Bid By", "First Name", "Last Name",
        "Email", "Phone", "Address", "City", "State",
        "Contractor", "Quote Amount", "Source File",
        "Customer Contact", "Customer Business", "Customer Address",
    ])
    for c in contacts:
        csi_val = extract_csi_from_filename(c.get("source_filename", ""))
        writer.writerow([
            csi_val,
            c.get("sub_contractor", ""), c.get("bid_by", ""),
            c.get("first_name", ""), c.get("last_name", ""),
            c.get("email", ""), c.get("phone", ""),
            c.get("address", ""), c.get("city", ""), c.get("state", ""),
            c.get("contractor", ""), c.get("quote_amount", ""),
            c.get("source_filename", ""),
            c.get("customer_contact_name", ""), c.get("customer_business", ""), c.get("customer_address", ""),
        ])
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=contacts_{run_id[:8]}.csv"}
    )

@api_router.get("/runs/{run_id}/download/errors")
async def download_errors_csv(run_id: str, request: Request):
    user = await get_current_user(request)
    errors = await db.processing_errors.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).to_list(1000)
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Filename", "Reason", "Missing Fields"])
    for e in errors:
        writer.writerow([e.get("filename",""), e.get("reason",""), e.get("missing_fields","")])
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=error_report_{run_id[:8]}.csv"}
    )

# =============================================================================
# FILE PROCESSING LOG — JSON view + XLSX download
# =============================================================================
LOG_COLUMNS = [
    ("filename",          "Filename",            40),
    ("file_type",         "File Type",           12),
    ("size_kb",           "File Size (KB)",      12),
    ("sha256",            "SHA-256 (content hash)", 18),
    ("csi",               "CSI",                  6),
    ("status",            "Status",              18),
    ("contacts_extracted","Contacts Extracted",  14),
    ("processing_tool",   "Processing Tool",     34),
    ("llm_model",         "LLM Model",           18),
    ("llm_provider",      "LLM Provider",        22),
    ("support_tools",     "Support Tools",       34),
    ("pages_sent",        "Pages Sent to LLM",   12),
    ("started_at",        "Started At",          22),
    ("completed_at",      "Completed At",        22),
    ("duration_sec",      "Duration (sec)",      12),
    ("missing_fields",    "Missing Fields",      30),
    ("issue_reason",      "Issue / Error Reason",50),
    ("retries",           "Retries",              8),
]

@api_router.get("/runs/{run_id}/log")
async def get_run_log(run_id: str, request: Request):
    """JSON view of the file-processing log for this run (live). Falls back
    to a synthesized log from files+contacts+errors if file_logs is empty."""
    user = await get_current_user(request)
    rows = await db.file_logs.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).sort("filename", 1).to_list(50000)
    if rows:
        return rows
    # Synthesize (mirrors the XLSX fallback)
    files = await db.files.find({"run_id": run_id, "user_id": user["_id"]},
                                 {"_id": 0, "original_filename": 1, "size": 1, "sha256": 1, "created_at": 1}
                                ).sort("original_filename", 1).to_list(50000)
    contacts_per_file = {}
    async for row in db.contacts.aggregate([
        {"$match": {"run_id": run_id, "user_id": user["_id"]}},
        {"$group": {"_id": "$source_filename", "n": {"$sum": 1}}}
    ]):
        contacts_per_file[row.get("_id", "")] = row.get("n", 0)
    errors_by_file = {}
    async for e in db.processing_errors.find(
        {"run_id": run_id, "user_id": user["_id"]},
        {"_id": 0, "filename": 1, "reason": 1, "missing_fields": 1}
    ):
        errors_by_file[e.get("filename", "")] = e
    out = []
    for f in files:
        fname = f.get("original_filename", "")
        ext = fname.lower().rsplit(".", 1)[-1] if "." in fname else ""
        contacts_n = contacts_per_file.get(fname, 0)
        err = errors_by_file.get(fname)
        if err:
            status, issue, missing = "Failure", err.get("reason", ""), err.get("missing_fields", "")
        elif contacts_n > 0:
            status, issue, missing = "Success", "", ""
        else:
            status, issue, missing = "No Contacts Found", "No contacts returned (log row pruned — reconstructed)", ""
        out.append({
            "filename": fname, "file_type": FILE_TYPE_MAP.get(ext, ext.upper() or "Unknown"),
            "size_kb": round(f.get("size", 0) / 1024), "sha256": f.get("sha256", ""),
            "csi": extract_csi_from_filename(fname),
            "status": status, "contacts_extracted": contacts_n,
            "processing_tool": "(retention: log row pruned — reconstructed)",
            "llm_model": "gemini-2.5-flash", "llm_provider": "Google (Emergent Key)",
            "support_tools": "", "pages_sent": 0,
            "started_at": "", "completed_at": f.get("created_at", ""),
            "duration_sec": 0, "missing_fields": missing, "issue_reason": issue, "retries": 0,
        })
    return out

@api_router.get("/runs/{run_id}/download/log")
async def download_run_log_xlsx(run_id: str, request: Request):
    """Download the file-processing log as a styled XLSX workbook.
    If file_logs rows don't exist for this run (e.g., older runs predating the
    log, or legacy data), synthesize rows on the fly from files + contacts +
    processing_errors so the download always reflects the true status."""
    user = await get_current_user(request)
    # Verify run belongs to user
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0, "created_at": 1})
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    rows = await db.file_logs.find({"run_id": run_id, "user_id": user["_id"]}, {"_id": 0}).sort("filename", 1).to_list(50000)

    # Fallback synthesis — when file_logs is empty but the run has data
    if not rows:
        # Pull files for this run
        files = await db.files.find({"run_id": run_id, "user_id": user["_id"]},
                                     {"_id": 0, "original_filename": 1, "size": 1, "sha256": 1, "created_at": 1}
                                    ).sort("original_filename", 1).to_list(50000)
        # Contacts by filename
        contacts_per_file = {}
        async for row in db.contacts.aggregate([
            {"$match": {"run_id": run_id, "user_id": user["_id"]}},
            {"$group": {"_id": "$source_filename", "n": {"$sum": 1}}}
        ]):
            contacts_per_file[row.get("_id", "")] = row.get("n", 0)
        # Errors by filename
        errors_by_file = {}
        async for e in db.processing_errors.find(
            {"run_id": run_id, "user_id": user["_id"]},
            {"_id": 0, "filename": 1, "reason": 1, "missing_fields": 1}
        ):
            errors_by_file[e.get("filename", "")] = e
        # Duplicates by filename
        dupes_per_file = {}
        async for row in db.duplicates.aggregate([
            {"$match": {"run_id": run_id, "user_id": user["_id"]}},
            {"$group": {"_id": "$source_filename", "n": {"$sum": 1}}}
        ]):
            dupes_per_file[row.get("_id", "")] = row.get("n", 0)

        for f in files:
            fname = f.get("original_filename", "")
            ext = fname.lower().rsplit(".", 1)[-1] if "." in fname else ""
            contacts_n = contacts_per_file.get(fname, 0)
            dupes_n = dupes_per_file.get(fname, 0)
            err = errors_by_file.get(fname)
            if err:
                status = "Failure"
                issue = err.get("reason", "")
                missing = err.get("missing_fields", "")
            elif contacts_n > 0:
                status = "Success"; issue = ""; missing = ""
            elif dupes_n > 0:
                status = "Success"; issue = f"Produced {dupes_n} duplicate contact(s) that were merged into prior rows."; missing = ""
            else:
                status = "No Contacts Found"
                issue = "No contacts returned (reason not logged — file_log row missing)"
                missing = ""
            rows.append({
                "filename": fname,
                "file_type": FILE_TYPE_MAP.get(ext, ext.upper() or "Unknown"),
                "size_kb": round(f.get("size", 0) / 1024),
                "sha256": f.get("sha256", ""),
                "csi": extract_csi_from_filename(fname),
                "status": status,
                "contacts_extracted": contacts_n,
                "processing_tool": "(retention: log row pruned — reconstructed)",
                "llm_model": "gemini-2.5-flash",
                "llm_provider": "Google (Emergent Key)",
                "support_tools": "",
                "pages_sent": 0,
                "started_at": "", "completed_at": f.get("created_at", ""),
                "duration_sec": 0,
                "missing_fields": missing,
                "issue_reason": issue,
                "retries": 0,
            })

    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Processing Log"

    # Header row
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="1E293B")
    for idx, (_, label, width) in enumerate(LOG_COLUMNS, start=1):
        cell = ws.cell(row=1, column=idx, value=label)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(LOG_COLUMNS))}1"

    # Status color map
    status_fills = {
        "Success":            PatternFill("solid", fgColor="D1FAE5"),  # green tint
        "Partial Success":    PatternFill("solid", fgColor="FEF3C7"),  # yellow tint
        "No Contacts Found":  PatternFill("solid", fgColor="E0E7FF"),  # blue tint
        "Failure":            PatternFill("solid", fgColor="FEE2E2"),  # red tint
        "Skipped":            PatternFill("solid", fgColor="F1F5F9"),  # gray tint
        "Skipped (Duplicate Content)": PatternFill("solid", fgColor="EDE9FE"),  # purple tint
        "In Progress":        PatternFill("solid", fgColor="CFFAFE"),  # cyan tint
        "Queued":             PatternFill("solid", fgColor="F8FAFC"),  # very light gray
    }

    for r_idx, row in enumerate(rows, start=2):
        for c_idx, (key, _, _) in enumerate(LOG_COLUMNS, start=1):
            val = row.get(key, "")
            ws.cell(row=r_idx, column=c_idx, value=val)
        # Color the Status column (5th)
        status_val = row.get("status", "")
        fill = status_fills.get(status_val)
        if fill:
            ws.cell(row=r_idx, column=5).fill = fill
            ws.cell(row=r_idx, column=5).font = Font(bold=True)

    # Summary block at bottom
    summary_row = len(rows) + 3
    status_counts = {}
    for row in rows:
        s = row.get("status", "Unknown")
        status_counts[s] = status_counts.get(s, 0) + 1
    ws.cell(row=summary_row, column=1, value="SUMMARY").font = Font(bold=True, size=12)
    summary_row += 1
    ws.cell(row=summary_row, column=1, value="Total Files").font = Font(bold=True)
    ws.cell(row=summary_row, column=2, value=len(rows))
    for status, count in sorted(status_counts.items()):
        summary_row += 1
        ws.cell(row=summary_row, column=1, value=status).font = Font(bold=True)
        ws.cell(row=summary_row, column=2, value=count)
        fill = status_fills.get(status)
        if fill:
            ws.cell(row=summary_row, column=1).fill = fill

    # Stream out
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename=processing_log_{run_id[:8]}.xlsx"}
    )

# =============================================================================
# SKIP REGISTRY — portable cross-run dedup list. Export before the container
# shuts down, re-import after rebuild so the system still knows which files
# have already been analyzed.
# =============================================================================
@api_router.get("/skip-registry/export")
async def export_skip_registry(request: Request):
    """Download a CSV of every file this user has processed, keyed by sha256 when
    available and by (filename, size) otherwise. Re-import to restore dedup after
    a container rebuild.
    """
    user = await get_current_user(request)

    # Pre-count contacts per (run_id, source_filename) so we can attach counts
    # to files that predate the file_logs collection. Single aggregation query.
    contacts_counts = {}  # (run_id, filename) -> count
    agg = db.contacts.aggregate([
        {"$match": {"user_id": user["_id"]}},
        {"$group": {"_id": {"run_id": "$run_id", "fn": "$source_filename"}, "n": {"$sum": 1}}}
    ])
    async for row in agg:
        k = (row["_id"].get("run_id", ""), row["_id"].get("fn", ""))
        contacts_counts[k] = row["n"]

    files_cursor = db.files.find(
        {"user_id": user["_id"]},
        {"_id": 0, "sha256": 1, "original_filename": 1, "size": 1, "run_id": 1, "created_at": 1}
    )
    seen_hashes = set()
    seen_legacy = set()
    rows = []
    async for f in files_cursor:
        h = f.get("sha256") or ""
        fname = f.get("original_filename", "")
        size = f.get("size", 0)
        if h:
            if h in seen_hashes:
                continue
            seen_hashes.add(h)
        else:
            key = f"LEGACY::{fname}::{size}"
            if key in seen_legacy:
                continue
            seen_legacy.add(key)
        # Prefer log row for status, fall back to aggregated contacts count
        log = await db.file_logs.find_one(
            {"user_id": user["_id"], "run_id": f.get("run_id"), "filename": fname},
            {"_id": 0, "status": 1, "contacts_extracted": 1, "completed_at": 1}
        )
        if not log and h:
            log = await db.file_logs.find_one(
                {"user_id": user["_id"], "sha256": h},
                {"_id": 0, "status": 1, "contacts_extracted": 1, "completed_at": 1}
            )
        # Backfill contacts_count from contacts collection when log is missing/empty
        cc = (log or {}).get("contacts_extracted", 0)
        if not cc:
            cc = contacts_counts.get((f.get("run_id", ""), fname), 0)
        # Backfill status similarly
        status = (log or {}).get("status") or ("Success" if cc > 0 else "Processed")
        rows.append({
            "filename": fname,
            "sha256": h,
            "size": size,
            "processed_at": (log or {}).get("completed_at") or f.get("created_at", ""),
            "run_id": f.get("run_id", ""),
            "contacts_count": cc,
            "status": status,
        })

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["filename", "sha256", "size", "processed_at", "run_id", "contacts_count", "status"])
    for r in rows:
        writer.writerow([r["filename"], r["sha256"], r["size"], r["processed_at"], r["run_id"], r["contacts_count"], r["status"]])
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=skip_registry_{datetime.now(timezone.utc).strftime('%Y%m%d')}.csv"}
    )


@api_router.post("/skip-registry/import")
async def import_skip_registry(request: Request, file: UploadFile = File(...)):
    """Re-populate the dedup registry from a previously exported CSV.
    Creates minimal `files` records so future uploads with matching
    sha256 (or filename+size for legacy entries) are skipped.
    Safe to re-import repeatedly (idempotent).
    """
    user = await get_current_user(request)
    if not (file.filename or "").lower().endswith(".csv"):
        raise HTTPException(status_code=400, detail="Expected a CSV file exported from /skip-registry/export")
    data = await file.read()
    try:
        text = data.decode("utf-8", errors="replace")
    except Exception:
        raise HTTPException(status_code=400, detail="Unreadable file")
    reader = csv.DictReader(io.StringIO(text))
    if "filename" not in (reader.fieldnames or []):
        raise HTTPException(status_code=400, detail="CSV missing required column 'filename'")
    imported_run_id = f"imported-{uuid.uuid4()}"
    imported = 0
    skipped = 0
    for row in reader:
        sha = (row.get("sha256") or "").strip()
        fname = (row.get("filename") or "").strip()
        try:
            size = int(row.get("size") or 0)
        except ValueError:
            size = 0
        if not fname:
            skipped += 1
            continue
        # Idempotent check — sha first, then legacy (filename+size)
        existing = None
        if sha and len(sha) == 64:
            existing = await db.files.find_one({"user_id": user["_id"], "sha256": sha}, {"_id": 0, "id": 1})
        if not existing:
            existing = await db.files.find_one(
                {"user_id": user["_id"], "original_filename": fname, "size": size},
                {"_id": 0, "id": 1}
            )
        if existing:
            skipped += 1
            continue
        await db.files.insert_one({
            "id": str(uuid.uuid4()),
            "run_id": row.get("run_id") or imported_run_id,
            "user_id": user["_id"],
            "storage_path": "",
            "original_filename": fname,
            "content_type": "",
            "size": size,
            "sha256": sha if (sha and len(sha) == 64) else "",
            "is_deleted": True,
            "imported_from_registry": True,
            "imported_processed_at": row.get("processed_at") or "",
            "imported_contacts_count": int(row.get("contacts_count") or 0),
            "imported_status": row.get("status") or "",
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        imported += 1
    return {"imported": imported, "skipped_already_known": skipped, "total_in_file": imported + skipped}

# =============================================================================
# MASTER INDEX — user uploads a CSV/XLSX with a FileName column; app compares
# against the files/contacts/errors collections to report which files have
# been processed vs. still need processing. Stored per-user (replace on re-upload).
# =============================================================================
def _parse_master_index(filename: str, data: bytes) -> List[str]:
    """Return a list of filenames parsed from a CSV or XLSX. Column 'FileName'
    (case-insensitive) preferred; if no header or a single-column sheet, the
    first column is used. Filenames are kept case-sensitive as stored."""
    lower = (filename or "").lower()
    names: List[str] = []
    if lower.endswith(".xlsx"):
        from openpyxl import load_workbook
        try:
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Unreadable XLSX: {e}")
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            return []
        # Detect header row
        header_idx = None
        for i, r in enumerate(rows[:1]):
            for c in r:
                if isinstance(c, str) and c.strip().lower() in ("filename", "file name", "file_name"):
                    header_idx = i
                    break
            if header_idx is not None:
                break
        col_idx = 0
        if header_idx is not None:
            header_row = rows[header_idx]
            for j, c in enumerate(header_row):
                if isinstance(c, str) and c.strip().lower() in ("filename", "file name", "file_name"):
                    col_idx = j
                    break
            data_rows = rows[header_idx + 1:]
        else:
            data_rows = rows
        for r in data_rows:
            if col_idx < len(r) and r[col_idx] is not None:
                s = str(r[col_idx]).strip()
                if s:
                    names.append(s)
    elif lower.endswith(".csv"):
        try:
            text = data.decode("utf-8-sig", errors="replace")
        except Exception:
            raise HTTPException(status_code=400, detail="Unreadable CSV")
        # Try DictReader first — matches header 'FileName'/'filename'/'file name'/...
        reader = csv.reader(io.StringIO(text))
        all_rows = [r for r in reader if any((c or "").strip() for c in r)]
        if not all_rows:
            return []
        header = all_rows[0]
        col_idx = 0
        has_header = False
        for j, c in enumerate(header):
            if (c or "").strip().lower() in ("filename", "file name", "file_name"):
                col_idx = j
                has_header = True
                break
        data_rows = all_rows[1:] if has_header else all_rows
        for r in data_rows:
            if col_idx < len(r):
                s = (r[col_idx] or "").strip()
                if s:
                    names.append(s)
    else:
        raise HTTPException(status_code=400, detail="Expected a .csv or .xlsx file")
    return names


def _file_ext(fname: str) -> str:
    if not fname:
        return "(none)"
    i = fname.rfind(".")
    return fname[i:].lower() if i >= 0 else "(none)"


@api_router.post("/master-index/upload")
async def upload_master_index(request: Request, file: UploadFile = File(...)):
    """Upload a master index (CSV or XLSX, column 'FileName')."""
    user = await get_current_user(request)
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")
    names = _parse_master_index(file.filename or "", data)
    if not names:
        raise HTTPException(status_code=400, detail="No filenames found. Ensure the file has a 'FileName' column or a single column of filenames.")
    # Preserve order & dedupe case-sensitively
    seen = set()
    unique = []
    for n in names:
        if n not in seen:
            seen.add(n)
            unique.append(n)
    doc = {
        "user_id": user["_id"],
        "original_filename": file.filename or "master_index",
        "entries": unique,
        "total": len(unique),
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }
    await db.master_index.replace_one({"user_id": user["_id"]}, doc, upsert=True)
    return {"total": len(unique), "original_filename": doc["original_filename"], "uploaded_at": doc["uploaded_at"]}


async def _compare_master_index(user_id: str):
    """Compare the stored master index against files/contacts/errors.
    Returns (summary, results list).
    Status values:
      - "Processed"             — file exists AND produced >=1 contact
      - "Processed (no contacts)" — file exists, run completed, no contacts/errors
      - "Error"                 — file exists AND has a processing error
      - "Pending"               — file uploaded but run not yet completed
      - "Not Uploaded"          — filename never seen, OR the run that uploaded
                                  it was deleted (orphaned file rows)
    """
    master = await db.master_index.find_one({"user_id": user_id}, {"_id": 0})
    if not master:
        return None, None
    entries = master.get("entries", [])
    if not entries:
        return master, []

    lookups: dict = {name: {"has_file": False, "run_id": "", "size": 0, "processed_at": "",
                             "contacts_count": 0, "has_error": False, "run_status": "",
                             "archive_filename": None} for name in entries}

    async for f in db.files.find(
        {"user_id": user_id, "original_filename": {"$in": entries}},
        {"_id": 0, "original_filename": 1, "run_id": 1, "size": 1, "created_at": 1, "archive_filename": 1}
    ):
        nm = f.get("original_filename", "")
        row = lookups.get(nm)
        if row is not None:
            row["has_file"] = True
            if f.get("created_at", "") > row["processed_at"]:
                row["processed_at"] = f.get("created_at", "")
                row["run_id"] = f.get("run_id", "")
                row["size"] = f.get("size", 0)
                row["archive_filename"] = f.get("archive_filename")

    # Fetch statuses of involved runs. Any run_id that does NOT come back is
    # ORPHANED (run was deleted) and we treat the file as "Not Uploaded" so
    # the user can re-upload it; SHA-256 dedup will skip the LLM if bytes match.
    involved_run_ids = {r["run_id"] for r in lookups.values() if r["run_id"]}
    seen_runs = set()
    if involved_run_ids:
        async for r in db.runs.find(
            {"user_id": user_id, "id": {"$in": list(involved_run_ids)}},
            {"_id": 0, "id": 1, "status": 1}
        ):
            seen_runs.add(r["id"])
            for info in lookups.values():
                if info["run_id"] == r.get("id"):
                    info["run_status"] = r.get("status", "")
    # Mark orphans (run was deleted but file rows remain). Count them
    # separately so the UI can hint that they were reclassified.
    orphan_count = 0
    orphan_run_ids = involved_run_ids - seen_runs
    if orphan_run_ids:
        for info in lookups.values():
            if info["run_id"] in orphan_run_ids:
                info["has_file"] = False  # Treat as Not Uploaded
                info["run_id"] = ""
                info["run_status"] = ""
                orphan_count += 1

    # Contacts counts per filename
    agg = db.contacts.aggregate([
        {"$match": {"user_id": user_id, "source_filename": {"$in": entries}}},
        {"$group": {"_id": "$source_filename", "n": {"$sum": 1}}},
    ])
    async for row in agg:
        nm = row.get("_id", "")
        if nm in lookups:
            lookups[nm]["contacts_count"] = row.get("n", 0)

    err_names_cursor = db.processing_errors.find(
        {"user_id": user_id, "filename": {"$in": entries}},
        {"_id": 0, "filename": 1}
    )
    async for e in err_names_cursor:
        nm = e.get("filename", "")
        if nm in lookups:
            lookups[nm]["has_error"] = True

    results = []
    summary = {"total": len(entries), "processed": 0, "processed_no_contacts": 0,
               "errored": 0, "pending": 0, "not_uploaded": 0,
               "orphaned_reclassified": orphan_count}
    by_type: dict = {}
    run_completed_statuses = {"completed", "cancelled"}

    for nm in entries:
        info = lookups[nm]
        if not info["has_file"]:
            status = "Not Uploaded"
            summary["not_uploaded"] += 1
        elif info["contacts_count"] > 0:
            status = "Processed"
            summary["processed"] += 1
        elif info["has_error"]:
            status = "Error"
            summary["errored"] += 1
        elif info["run_status"] in run_completed_statuses:
            status = "Processed (no contacts)"
            summary["processed_no_contacts"] += 1
        else:
            status = "Pending"
            summary["pending"] += 1

        ext = _file_ext(nm)
        if ext not in by_type:
            by_type[ext] = {"total": 0, "processed": 0, "processed_no_contacts": 0,
                            "errored": 0, "pending": 0, "not_uploaded": 0}
        by_type[ext]["total"] += 1
        key = {"Processed": "processed", "Processed (no contacts)": "processed_no_contacts",
               "Error": "errored", "Pending": "pending", "Not Uploaded": "not_uploaded"}[status]
        by_type[ext][key] += 1

        results.append({
            "filename": nm,
            "extension": ext,
            "status": status,
            "run_id": info["run_id"],
            "run_status": info["run_status"],
            "size": info["size"],
            "processed_at": info["processed_at"],
            "contacts_count": info["contacts_count"],
            "archive_filename": info["archive_filename"],
        })

    summary["by_type"] = [
        {"extension": k, **v} for k, v in sorted(by_type.items(), key=lambda x: -x[1]["total"])
    ]
    return master, {"summary": summary, "results": results}


@api_router.get("/master-index")
async def get_master_index(request: Request):
    """Return the master index metadata + comparison results. Null if none uploaded."""
    user = await get_current_user(request)
    master, comparison = await _compare_master_index(user["_id"])
    if not master:
        return {"master": None, "comparison": None}
    return {
        "master": {
            "original_filename": master.get("original_filename", ""),
            "total": master.get("total", 0),
            "uploaded_at": master.get("uploaded_at", ""),
        },
        "comparison": comparison,
    }


@api_router.get("/master-index/download")
async def download_master_index_comparison(request: Request):
    """Download the comparison results as a CSV."""
    user = await get_current_user(request)
    master, comparison = await _compare_master_index(user["_id"])
    if not master or not comparison:
        raise HTTPException(status_code=404, detail="No master index uploaded")
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["FileName", "Extension", "Status", "Contacts Extracted", "Source Archive", "Run ID", "Run Status", "Size", "Processed At"])
    for r in comparison["results"]:
        writer.writerow([r["filename"], r["extension"], r["status"], r["contacts_count"], r.get("archive_filename") or "", r["run_id"], r.get("run_status",""), r["size"], r["processed_at"]])
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=master_index_comparison_{datetime.now(timezone.utc).strftime('%Y%m%d')}.csv"}
    )


@api_router.delete("/master-index")
async def delete_master_index(request: Request):
    user = await get_current_user(request)
    await db.master_index.delete_many({"user_id": user["_id"]})
    return {"message": "Master index cleared"}


# =============================================================================
# DELETE ROUTES
# =============================================================================
@api_router.delete("/runs/{run_id}")
async def delete_run(run_id: str, request: Request):
    user = await get_current_user(request)
    run = await db.runs.find_one({"id": run_id, "user_id": user["_id"]}, {"_id": 0})
    if not run:
        raise HTTPException(status_code=404, detail="Run not found")
    await db.runs.delete_one({"id": run_id, "user_id": user["_id"]})
    await db.contacts.delete_many({"run_id": run_id, "user_id": user["_id"]})
    await db.processing_errors.delete_many({"run_id": run_id, "user_id": user["_id"]})
    await db.duplicates.delete_many({"run_id": run_id, "user_id": user["_id"]})
    await db.file_logs.delete_many({"run_id": run_id, "user_id": user["_id"]})
    await db.progress.delete_many({"run_id": run_id})
    await db.files.update_many({"run_id": run_id, "user_id": user["_id"]}, {"$set": {"is_deleted": True}})
    return {"message": "Run deleted"}

@api_router.delete("/data/all")
async def delete_all_data(request: Request):
    user = await get_current_user(request)
    uid = user["_id"]
    results = {}
    results["runs"] = (await db.runs.delete_many({"user_id": uid})).deleted_count
    results["contacts"] = (await db.contacts.delete_many({"user_id": uid})).deleted_count
    results["errors"] = (await db.processing_errors.delete_many({"user_id": uid})).deleted_count
    results["duplicates"] = (await db.duplicates.delete_many({"user_id": uid})).deleted_count
    results["file_logs"] = (await db.file_logs.delete_many({"user_id": uid})).deleted_count
    results["files"] = (await db.files.update_many({"user_id": uid}, {"$set": {"is_deleted": True}})).modified_count
    run_ids = [r["run_id"] async for r in db.progress.find({})]
    await db.progress.delete_many({})
    return {"message": "All data deleted", "deleted": results}

# =============================================================================
# STARTUP / MIDDLEWARE
# =============================================================================
@app.on_event("startup")
async def startup():
    # Ensure poppler is installed (needed for pdf2image)
    import subprocess
    try:
        subprocess.run(["which", "pdftoppm"], check=True, capture_output=True)
    except subprocess.CalledProcessError:
        logger.warning("poppler-utils not found, installing...")
        subprocess.run(["apt-get", "install", "-y", "poppler-utils"], capture_output=True)
    # Ensure LibreOffice is installed (needed for DOCX/XLSX vision fallback)
    try:
        subprocess.run(["which", "libreoffice"], check=True, capture_output=True)
    except subprocess.CalledProcessError:
        logger.warning("libreoffice not found, installing in background...")
        subprocess.Popen(["apt-get", "install", "-y", "--no-install-recommends", "libreoffice"],
                         stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    try:
        init_storage()
        logger.info("Object storage initialized")
    except Exception as e:
        logger.error(f"Storage init failed: {e}")
    await db.users.create_index("email", unique=True)
    await db.login_attempts.create_index("identifier")
    await db.runs.create_index([("user_id", 1), ("created_at", -1)])
    await db.contacts.create_index("run_id")
    await db.processing_errors.create_index("run_id")
    await db.files.create_index("run_id")
    await db.progress.create_index("run_id")
    await db.file_logs.create_index([("run_id", 1), ("filename", 1)])
    await db.file_logs.create_index([("user_id", 1), ("run_id", 1)])
    await db.files.create_index([("user_id", 1), ("sha256", 1)])
    await db.files.create_index([("user_id", 1), ("original_filename", 1), ("size", 1)])
    await db.duplicates.create_index("run_id")
    await db.raw_contacts.create_index("run_id")
    await db.master_index.create_index("user_id", unique=True)
    # Recover stale "processing" runs from server restarts
    stale = await db.runs.count_documents({"status": "processing"})
    if stale > 0:
        await db.runs.update_many({"status": "processing"}, {"$set": {"status": "stale"}})
        await db.progress.update_many({"status": "processing"}, {"$set": {"status": "stale", "message": "Processing interrupted — click Retry to resume"}})
        logger.warning(f"Marked {stale} stale processing runs for retry")
    # One-time self-heal: recompute stats for every completed run so historical
    # stats drift (from old retention / bad formulas) is corrected.
    completed = await db.runs.find({"status": "completed"}, {"_id": 0, "id": 1, "user_id": 1}).to_list(5000)
    healed = 0
    for r in completed:
        try:
            await recompute_run_stats(r["id"], r["user_id"])
            healed += 1
        except Exception as e:
            logger.warning(f"Stats recompute failed for {r.get('id','?')}: {e}")
    if healed:
        logger.info(f"Startup stats self-heal: recomputed {healed} completed run(s)")
    # Seed admin
    admin_email = os.environ.get("ADMIN_EMAIL", "admin@trueflow.com")
    admin_password = os.environ.get("ADMIN_PASSWORD", "TrueFlow2024!")
    existing = await db.users.find_one({"email": admin_email})
    if not existing:
        hashed = hash_password(admin_password)
        result = await db.users.insert_one({
            "email": admin_email, "password_hash": hashed,
            "name": "Admin", "role": "admin",
            "must_change_password": True,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        user_id = str(result.inserted_id)
        await db.settings.insert_one({
            "user_id": user_id,
            "exclusion_domain": "horizonc.com",
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        logger.info(f"Admin user created: {admin_email}")
    elif not verify_password(admin_password, existing["password_hash"]):
        await db.users.update_one({"email": admin_email}, {"$set": {"password_hash": hash_password(admin_password), "must_change_password": True}})
    # Seed admin config if not exists
    admin_config = await db.admin_config.find_one({"key": "global"})
    if not admin_config:
        await db.admin_config.insert_one({
            "key": "global",
            "ai_model": "claude-sonnet",
            "claude_api_key": "",
            "openai_api_key": "",
            "max_pdfs_per_upload": 7000,
            "storage_max_mb": 750,
            "storage_target_mb": 300,
            "created_at": datetime.now(timezone.utc).isoformat()
        })
        logger.info("Admin config seeded")
    # Write test credentials
    os.makedirs("/app/memory", exist_ok=True)
    with open("/app/memory/test_credentials.md", "w") as f:
        f.write(f"# Test Credentials\n\n## Admin Account\n- Email: {admin_email}\n- Password: {admin_password}\n- Role: admin\n\n")
        f.write("## Auth Endpoints\n- POST /api/auth/register\n- POST /api/auth/login\n- POST /api/auth/logout\n- GET /api/auth/me\n- POST /api/auth/refresh\n")

app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[os.environ.get("FRONTEND_URL", "http://localhost:3000")],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown():
    client.close()
